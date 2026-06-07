import streamlit as st
import streamlit.components.v1 as components
import fitz  # PyMuPDF
import re
import json
import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER
try:
    from ehcp_parser.pipeline import EHCPPipeline
    EHCP_PARSER_AVAILABLE = True
except Exception as e:
    EHCP_PARSER_AVAILABLE = False
    print(f"EHCP parser unavailable: {e}")
    # ── SUPABASE CONNECTION ──────────────────────────────────────────────────────
try:
    from supabase import create_client, Client
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    SUPABASE_AVAILABLE = True
    print("Supabase connected successfully")
except Exception as e:
    SUPABASE_AVAILABLE = False
    supabase = None
    print(f"Supabase unavailable: {e}")
try:
    from sandbox.provision_inventory import build_provision_inventory
    INVENTORY_AVAILABLE = True

except Exception as e:
    INVENTORY_AVAILABLE = False
    print(f"Provision inventory unavailable: {e}")
# ── PAGE CONFIG ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="FRED — Families' Rights and Entitlements Directory",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── COLOUR CONSTANTS ─────────────────────────────────────────────────────────

RED    = "#C0392B"
AMBER  = "#D4A017"
GREEN  = "#1E8449"
NAVY   = "#1a2744"
LIGHT  = "#f7f9fc"

GOOGLE_FORM_URL = "https://docs.google.com/forms/d/e/1FAIpQLSeA1F9nEdQWkmplbAh973XKq2EsW0bEkhJiw7drhP7BZaPjKQ/viewform"

# ── CUSTOM CSS ───────────────────────────────────────────────────────────────

st.markdown(f"""
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500&display=swap');

  html, body, [class*="css"] {{
    font-family: 'DM Sans', sans-serif;
    background-color: #f2f0eb;
    color: #1a1a2e;
  }}

  h1, h2, h3 {{
    font-family: 'DM Serif Display', serif;
    font-weight: 400;
  }}

  /* Hero */
  .hero {{
    background: #2d4a2d;
    color: white;
    padding: 3rem 2.5rem 2.8rem;
    border-radius: 16px;
    text-align: center;
    margin-bottom: 2rem;
    position: relative;
    overflow: hidden;
  }}

  .hero::before {{
    content: '';
    position: absolute;
    top: -60px; right: -60px;
    width: 220px; height: 220px;
    border-radius: 50%;
    background: rgba(160,210,130,0.1);
  }}

  .hero h1 {{
    font-family: 'DM Serif Display', serif;
    font-size: 4rem;
    margin: 0;
    color: #e8f5e0;
    letter-spacing: 0.06em;
    font-weight: 400;
  }}

  .hero .full-name {{
    font-family: 'DM Sans', sans-serif;
    font-size: 0.75rem;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    color: #9dc98a;
    margin: 0.4rem 0 1rem 0;
    font-weight: 300;
  }}

  .hero .tagline {{
    font-family: 'DM Serif Display', serif;
    font-size: 1.45rem;
    font-weight: 400;
    color: #e8f5e0;
    margin-bottom: 0.5rem;
    line-height: 1.35;
  }}

  .hero .origin {{
    font-style: italic;
    color: #7ab870;
    margin-bottom: 1.4rem;
    font-size: 0.93rem;
    font-weight: 300;
  }}

  .hero .sub {{
    color: #b8d9a8;
    max-width: 540px;
    margin: 0 auto 1.4rem auto;
    font-size: 1rem;
    line-height: 1.7;
    font-weight: 300;
  }}

  .hero .service-line {{
    color: #7ab870;
    font-size: 0.82rem;
    margin-bottom: 1.8rem;
    letter-spacing: 0.04em;
  }}

  /* Gold CTA button */
  .cta-button {{
    display: inline-block;
    background: #e8c96a;
    color: #2d3a1a !important;
    font-family: 'DM Sans', sans-serif;
    font-weight: 600;
    font-size: 1rem;
    padding: 0.8rem 2.2rem;
    border-radius: 40px;
    text-decoration: none;
    letter-spacing: 0.02em;
    cursor: pointer;
    border: none;
  }}

  .reassurance {{
    color: #9dc98a;
    font-size: 0.85rem;
    margin-top: 1rem;
    font-weight: 300;
  }}

  .pricing-hint {{
    color: #7a8fa8;
    font-size: 0.82rem;
    margin-top: 0.4rem;
  }}

  /* Traffic light badges */
  .badge-red {{
    background: {RED};
    color: white;
    padding: 0.22rem 0.7rem;
    border-radius: 20px;
    font-weight: 500;
    font-size: 0.78rem;
    display: inline-block;
    margin-bottom: 0.5rem;
    letter-spacing: 0.03em;
  }}
  .badge-amber {{
    background: {AMBER};
    color: white;
    padding: 0.22rem 0.7rem;
    border-radius: 20px;
    font-weight: 500;
    font-size: 0.78rem;
    display: inline-block;
    margin-bottom: 0.5rem;
    letter-spacing: 0.03em;
  }}
  .badge-green {{
    background: {GREEN};
    color: white;
    padding: 0.22rem 0.7rem;
    border-radius: 20px;
    font-weight: 500;
    font-size: 0.78rem;
    display: inline-block;
    margin-bottom: 0.5rem;
    letter-spacing: 0.03em;
  }}

  /* Finding cards */
  .finding-red {{
    border-left: 4px solid {RED};
    background: #fdf4f3;
    padding: 1rem 1.2rem;
    border-radius: 0 8px 8px 0;
    margin-bottom: 1rem;
  }}
  .finding-amber {{
    border-left: 4px solid {AMBER};
    background: #fdf9f0;
    padding: 1rem 1.2rem;
    border-radius: 0 8px 8px 0;
    margin-bottom: 1rem;
  }}
  .finding-green {{
    border-left: 4px solid {GREEN};
    background: #f3faf5;
    padding: 1rem 1.2rem;
    border-radius: 0 8px 8px 0;
    margin-bottom: 1rem;
  }}

  /* How it works */
  .hiw-item {{
    display: flex;
    align-items: flex-start;
    gap: 1rem;
    margin-bottom: 0.8rem;
    padding: 0.9rem 1rem;
    background: white;
    border-radius: 10px;
    border: 0.5px solid #dde8dd;
  }}
  .hiw-num {{
    font-family: 'DM Serif Display', serif;
    font-size: 1.3rem;
    color: #7ab870;
    flex-shrink: 0;
    line-height: 1;
    width: 24px;
  }}

  /* Pricing cards */
  .pricing-card {{
    background: white;
    border: 0.5px solid #d0ddd0;
    border-radius: 12px;
    padding: 1.4rem 1.2rem;
    text-align: center;
    height: 100%;
  }}
  .pricing-card.featured {{
    border: 1.5px solid #7ab870;
  }}
  .pricing-card .price {{
    font-family: 'DM Serif Display', serif;
    font-size: 2rem;
    font-weight: 400;
    color: #2d4a2d;
  }}
  .pricing-card .period {{
    font-size: 0.82rem;
    color: #888;
    font-weight: 300;
  }}
  .pricing-card ul {{
    text-align: left;
    padding-left: 0;
    margin-top: 1rem;
    font-size: 0.88rem;
    line-height: 2;
    list-style: none;
    font-weight: 300;
    color: #555;
  }}
  .pricing-card ul li::before {{
    content: '·  ';
    color: #7ab870;
    font-weight: 700;
  }}

  .best-value-tag {{
    background: #2d4a2d;
    color: #e8f5e0;
    font-size: 0.7rem;
    font-weight: 500;
    padding: 0.2rem 0.7rem;
    border-radius: 20px;
    display: inline-block;
    margin-bottom: 0.5rem;
    letter-spacing: 0.08em;
    text-transform: uppercase;
  }}

  /* Sneak peek */
  .sneak-box {{
    border: 1.5px dashed #7ab870;
    border-radius: 12px;
    padding: 1.5rem;
    background: white;
    margin: 1.5rem 0;
  }}

  /* Document pills */
  .doc-pill {{
    display: inline-block;
    background: #eaf5e0;
    color: #2d5a2d;
    font-size: 0.78rem;
    padding: 0.28rem 0.8rem;
    border-radius: 20px;
    margin: 0.2rem;
    font-weight: 400;
  }}

  .section-divider {{
    border: none;
    border-top: 0.5px solid #dde8dd;
    margin: 2.5rem 0;
  }}

  .section-label {{
    font-size: 0.7rem;
    font-weight: 500;
    text-transform: uppercase;
    letter-spacing: 0.12em;
    color: #5a8a5a;
    margin-bottom: 0.6rem;
  }}

  /* Streamlit button override */
  .stButton > button {{
    background: #2d4a2d;
    color: white;
    border: none;
    border-radius: 40px;
    font-family: 'DM Sans', sans-serif;
    font-weight: 500;
    padding: 0.6rem 1.6rem;
    font-size: 0.97rem;
  }}
  .stButton > button:hover {{
    background: #3d5e3d;
    color: white;
  }}

  /* Primary CTA Streamlit button — gold */
  div[data-testid="stButton"] button[kind="primary"] {{
    background: #e8c96a;
    color: #2d3a1a;
    border-radius: 40px;
  }}

  footer {{visibility: hidden;}}
  #MainMenu {{visibility: hidden;}}

  [data-testid="stTextInput"] input {{
    border: 1.5px solid #9ab0c8 !important;
    background-color: #ffffff !important;
    border-radius: 6px !important;
  }}
</style>
""", unsafe_allow_html=True)

# ── EMAILJS ──────────────────────────────────────────────────────────────────


# ── PDF PARSER ────────────────────────────────────────────────────────────────
# Warwickshire format: multiple Section F blocks, two-column tables.
# Strategy: collect ALL text, identify section boundaries, extract each F block.

# ── VAGUE LANGUAGE — SECTION F LEGAL STANDARD ─────────────────────────────────
# Source: SOS!SEN EHCP Infosheet 2022, CoP para 9.69, Children and Families Act 2014
# Split into two groups so findings can distinguish commitment failures from
# specificity failures. Both are RED — both remove enforceability.

# Group 1: Modal language — weakens the commitment from "will" to conditional.
# The legal standard requires "will receive". Anything weaker is not a lawful
# commitment and cannot be enforced via Judicial Review.
VAGUE_MODAL = [
    "will benefit from",
    "would benefit from",
    "might benefit from",
    "should receive",
    "could receive",
    "may receive",
    "may be provided",
    "should be provided",
    "may be helpful",
    "might be appropriate",
    "would be appropriate",
    "should have access",
    "could have access",
]

# Group 2: Vague qualifiers — provision stated without specificity.
# SOS!SEN explicit list: "access to", "opportunities are", "regular", "up to",
# "as advised", "as required", "may be helpful", "contacts", "adults".
# Each one makes provision impossible to audit or enforce.
VAGUE_QUALIFIER = [
    "access to",
    "opportunities",
    "as advised",
    "as required",
    "as appropriate",
    "as needed",
    "where necessary",
    "up to",
    "at their discretion",
    "where possible",
    "as directed",
    "regular",
    "regularly",
    "ongoing support",
]

# ── TASK 2.3: EP / SALT / OT RECOMMENDATION TRIGGERS ────────────────────────
RECOMMENDATION_TRIGGERS = [
    "it is recommended that",
    "it is recommended",
    "would benefit from",
    "should receive",
    "recommended that",
    "recommended",
]

_REC_STOPWORDS = {
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "is", "are", "was", "were", "be", "been",
    "being", "have", "has", "had", "do", "does", "did", "will", "would",
    "could", "should", "may", "might", "shall", "can", "that", "this",
    "these", "those", "it", "its", "they", "them", "their", "he", "she",
    "his", "her", "we", "our", "you", "your", "i", "my", "who", "which",
    "also", "both", "through", "during", "before", "after", "above",
    "between", "out", "off", "over", "under", "again", "further", "then",
    "once", "into", "onto", "upon", "when", "where", "while", "such",
    "child", "pupil", "student", "young", "person", "their", "given",
    "continue", "continued", "ongoing", "therefore", "however","school",
    "staff", "able", "support", "needs", "help", "time",
    "work", "make", "feel", "find", "take", "give", "show", "know",
    "including", "within", "around", "across", "often", "always",
    "sometimes", "currently","sometimes", "currently", "specific", "having", 
    "benefit", "continue", "engage",
}

_REC_MATCH_THRESHOLD = 2

# Combined list — used where a single list is needed elsewhere in the app
PROHIBITED_WORDS = VAGUE_MODAL + VAGUE_QUALIFIER
# ── PROVISION LIBRARY ─────────────────────────────────────────────────────────
# Four categories matching the four need areas in analyse_section_b.
# Used by: analyse_section_b (B→F matching), build_provision_inventory,
# and the correspondence Operational Language Signal pattern.
# Duplicates removed. Source: FRED_PROVISION_LIBRARY.txt

PROVISION_LIBRARY = {
    "Communication and Interaction": [
        "visual supports", "now and next board", "communication passport",
        "comic strip conversations", "language intervention",
        "speech and language input", "social communication support",
        "processing time", "reduced language load", "makaton signing",
        "symbols", "structured language", "peer modelling", "rephrasing",
        "modified language", "augmented communication device",
        "clarified instructions", "visual sequencing", "task board",
        "communication prompts", "visual cue card", "language scaffolding",
        "structured conversation", "conversation prompts",
        "comprehension support", "attention cues", "adult language modelling",
        "supported interaction", "visual checklist", "task planner",
        "supported play", "transition warning", "structured transition",
        "adult facilitation", "social skills group", "peer support",
        "visual prompt", "attention support", "visual timetable",
        "now-and-next board", "social story", "comic strip conversation",
        "Colourful Semantics", "PECS", "AAC device", "Makaton signing",
        "pre-teaching vocabulary", "blank level questioning",
        "structured conversation starter", "communication book",
        "processing time pause", "LEGO therapy", "supported peer interaction",
        "explicit instruction", "first-then sequence",
        "concrete language model", "auditory processing breakdown",
        "visual prompt script", "talking partners structured framework",
        "word bank provision", "scaffolded sentence starter",
        "dual-coding delivery", "semantic mapping",
        "phonological awareness intervention", "cued articulation",
        "verbal prompt fading", "directed questioning scaffolding",
        "narrative therapy structure", "conversational turn-taking script",
        "minimal pair training", "total communication approach",
        "clarification tracking strategy", "literal language modification",
        "receptive language visual support", "expressive language framework",
        "echolalia functional redirection", "comprehension check constraint",
        "information carrying word framework",
    ],
    "Sensory and Physical": [
        "sensory diet", "sensory circuit", "movement opportunity",
        "movement opportunities", "movement break", "brain break",
        "fidget tool", "weighted item", "chewelry", "quiet workstation",
        "low arousal environment", "sensory regulation",
        "multi-sensory approach", "writing slope", "adapted resources",
        "prompt sheets", "supervision at unstructured times",
        "movement programme", "sensory integration support",
        "environmental adaptation", "alternative seating",
        "noise reduction support", "sensory toolkit", "regulation equipment",
        "fine motor support", "gross motor support", "adaptive writing tools",
        "postural support", "regulated movement", "calm workstation",
        "movement route", "sensory support plan", "ear defenders",
        "planned sensory break", "motor skills support",
        "environment adaptation", "low stimulus environment",
        "regulation support", "toileting support", "wobble cushion",
        "weighted lap pad", "noise-cancelling headphones",
        "slanted writing board", "pencil grip", "adaptive scissors",
        "sensory break area", "deep pressure input",
        "proprioceptive feedback task", "vestibular stimulation exercise",
        "tactile defensive modification", "environmental audit accommodation",
        "low-stimulus workspace", "privacy screen",
        "gross motor tracking program", "fine motor precision task",
        "handwriting grip stabilization", "spatial awareness grid",
        "visual tracking exercise", "sensory profile adaptation",
        "scheduled hydration prompt", "postural support seating",
        "soundfield system", "dimmed lighting provision",
        "tactile baseline object", "heavy work activity",
        "finger isolation exercise", "bilateral coordination task",
        "hand-eye coordination drill", "sensory decompression interval",
        "scheduled movement transition", "fidget regulation tool",
        "adaptive cutlery", "step-free route mapping",
        "physical prompt fading", "ergonomic workstation set-up",
    ],
    "SEMH": [
        "emotional regulation", "trusted adult", "check in", "safe adult",
        "time out card", "emotion coaching", "restorative conversation",
        "co-regulation", "regulation strategy", "anxiety support",
        "feelings fan", "safe space", "relationship support",
        "take-up time", "non-confrontational language", "emotional check in",
        "supported regulation", "planned regulation support",
        "emotion identification", "regulation breaks", "trusted relationship",
        "repair conversation", "emotional processing support",
        "adult reassurance", "predictable routine", "supported transition",
        "daily check in", "emotional support", "pastoral support",
        "regulation intervention", "exit card", "anxiety regulation",
        "Zones of Regulation", "5-point scale", "emotion thermometer",
        "safe space access", "de-escalation script", "box breathing technique",
        "predictable routine framing", "transition countdown",
        "object of reference", "emotional literacy checklist",
        "restorative conversation framework", "anxiety scaling tool",
        "regulation check-in schedule", "predictable exit strategy",
        "demand reduction protocol", "low-arousal approach",
        "safe adult check-in", "proprioceptive regulation task",
        "distraction technique blueprint", "hypervigilance reduction routine",
        "physiological regulation interval", "sensory grounding sequence",
        "structured playground role", "lunchtime nurture group",
        "calm-down toolkit", "self-regulation script",
        "contingency mapping visual", "reframing prompt tool",
        "structured check-out protocol", "emotional check-in board",
        "positive reinforcement schedule", "token economy framework",
        "antecedent modification sequence", "preferred activity transition card",
        "guided self-calming routine", "nurture room provision",
        "social script for conflict", "peer mediation model",
    ],
    "Cognition and Learning": [
        "task chunking", "scaffolded learning", "overlearning", "pre-teaching",
        "post teaching", "writing frame", "exam access arrangements",
        "processing support", "instruction breakdown",
        "differentiated resources", "chunking information", "small group work",
        "reading intervention", "numeracy intervention", "task reduction",
        "guided practice", "modelled learning", "supported recording",
        "visual task breakdown", "sequenced instructions",
        "working memory support", "repetition and reinforcement",
        "adult prompting", "independent learning support",
        "learning reinforcement", "structured recording support",
        "guided reading", "task support", "memory aids", "sequenced learning",
        "reduced worksheet", "alternative recording", "structured support",
        "learning prompts", "processing break", "overlearning routine",
        "precision teaching", "spaced retrieval practice",
        "dual-coding template", "word map scaffolding",
        "scaffolded worksheet", "dyscalculia targeted intervention",
        "phonics systematic recovery", "reading tracking ruler",
        "colored overlay use", "assistive technology text-to-speech",
        "speech-to-text software", "mind mapping framework",
        "visual checklist processing", "memory prompt cue card",
        "working memory externalization", "guided errorless learning",
        "multi-sensory spelling model", "backward chaining technique",
        "forward chaining technique", "diagnostic math baseline",
        "visual math manipulation tool", "Numicon framework",
        "Cuisenaire rod tracking", "graphic organizer template",
        "step-by-step instruction breakdown", "dictation method support",
        "exemplar reference framework", "task analysis breakdown",
        "high-frequency word bank", "cloze procedure scaffold",
        "enlarged print modification", "time management visual tracker",
        "independent work folder organization", "self-correction checklist",
        "modeled example framework", "faded scaffold sheet",
    ],
}

# Flat list for fast regex matching in correspondence analysis
PROVISION_TERMS_FLAT = [
    term for terms in PROVISION_LIBRARY.values() for term in terms
]

def _reconstruct_column(words):
    """
    Reconstruct readable text from pdfplumber word dicts.
    Groups words into lines by top coordinate, sorts by x within each line.
    Used to separate two-column PDF layouts.
    """
    if not words:
        return ""

    words_sorted = sorted(
        words,
        key=lambda w: (round(float(w['top']) / 5) * 5, float(w['x0']))
    )

    lines = []
    current_line = []
    current_top = None

    for word in words_sorted:
        word_top = round(float(word['top']) / 5) * 5
        if current_top is None or abs(word_top - current_top) <= 5:
            current_line.append(word)
            current_top = word_top
        else:
            if current_line:
                lines.append(sorted(current_line, key=lambda w: float(w['x0'])))
            current_line = [word]
            current_top = word_top

    if current_line:
        lines.append(sorted(current_line, key=lambda w: float(w['x0'])))

    text_lines = [" ".join(w['text'] for w in line) for line in lines if line]
    return "\n".join(text_lines)
def extract_text_from_pdf(uploaded_file):
    """
    Two-column aware PDF extraction.
    Uses pdfplumber for layout-aware extraction.
    Detects Warwickshire/Capita Synergy two-column Section F format
    and separates columns by x-coordinate to prevent text scrambling.
    Left column (provision text) is extracted cleanly first.
    Right column (deliverer info) is appended with a label.
    Falls back to PyMuPDF if pdfplumber fails.
    """
    try:
        import pdfplumber

        data = uploaded_file.read()
        pages_text = []

        with pdfplumber.open(BytesIO(data)) as pdf:
            page_count = len(pdf.pages)

            for page in pdf.pages:
                words = page.extract_words(
                    x_tolerance=3,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=True,
                )

                if not words:
                    pages_text.append("")
                    continue

                page_width = page.width

                # Column split threshold — Warwickshire format uses ~65% left / 35% right
                # A4 page = 595pt. Right column ("Who will provide support") starts ~390pt
                split_x = page_width * 0.60

                left_words  = [w for w in words if float(w['x0']) <= split_x]
                right_words = [w for w in words if float(w['x0']) >  split_x]

                # Only apply column separation if right column has meaningful content
                # Filters out page numbers and single-word headers
                right_content = [w for w in right_words if len(w['text']) > 2]
                is_two_column = (
                    len(right_content) >= 5 and
                    len(left_words)    >= 10
                )

                if is_two_column:
                    left_text  = _reconstruct_column(left_words)
                    right_text = _reconstruct_column(right_words)
                    if right_text.strip():
                        page_text = left_text + "\nDeliverer: " + right_text
                    else:
                        page_text = left_text
                else:
                    page_text = page.extract_text(
                        x_tolerance=3, y_tolerance=3
                    ) or ""

                pages_text.append(page_text)

        result = "\n".join(pages_text)
        print(f"pdfplumber extracted: {len(result)} chars from {page_count} pages")

        if len(result.strip()) < 50:
            return "__PDF_EXTRACTION_FAILED__"

        return result

    except Exception as e:
        print(f"pdfplumber error: {e} — falling back to PyMuPDF")
        try:
            uploaded_file.seek(0)
            data = uploaded_file.read()
            doc  = fitz.open(stream=data, filetype="pdf")
            pages = []
            for page in doc:
                text = page.get_text("text")
                if not text.strip():
                    blocks = page.get_text("blocks")
                    text = "\n".join(
                        b[4] for b in blocks if isinstance(b[4], str)
                    )
                if not text.strip():
                    d = page.get_text("dict")
                    words = []
                    for block in d.get("blocks", []):
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                words.append(span.get("text", ""))
                    text = " ".join(words)
                pages.append(text)
            result = "\n".join(pages)
            print(f"PyMuPDF fallback extracted: {len(result)} chars")
            return result
        except Exception as e2:
            print(f"PyMuPDF fallback error: {e2}")
            return ""

def extract_text_from_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""


def extract_text_from_txt(uploaded_file):
    try:
        return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


def extract_text(uploaded_file):
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        text = extract_text_from_pdf(uploaded_file)
        if len(text.strip()) < 50:
            return "__PDF_EXTRACTION_FAILED__"
        return text
    elif name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    elif name.endswith(".txt"):
        return extract_text_from_txt(uploaded_file)
    return ""

def find_section_blocks(full_text, section_letter):
    NEED_AREA_BOUNDARY = re.compile(
        r'\b(?:Communication\s+and\s+Interaction|'
        r'Social,?\s+Emotional\s+and\s+Mental\s+Health|'
        r'Cognition\s+and\s+Learning|'
        r'Physical\s+(?:and/or|and|or)\s+Sensory|'
        r'Health\s+Needs?|'
        r'Social\s+Care)\b',
        re.IGNORECASE
    )

    alt_pattern = re.compile(
        rf'\bSection\s+{section_letter}\b',
        re.IGNORECASE
    )

    positions = [m.start() for m in alt_pattern.finditer(full_text)]

    deduped = []
    for pos in sorted(positions):
        if not deduped or pos - deduped[-1] > 50:
            deduped.append(pos)

    blocks = []

    all_section_starts = sorted(
        [m.start() for m in re.finditer(
            r'\bSection\s+[A-Z]\b', full_text, re.IGNORECASE
        )] +
        [m.start() for m in NEED_AREA_BOUNDARY.finditer(full_text)]
    )

    for i, start in enumerate(deduped):
        next_starts = [s for s in all_section_starts if s > start + 10]
        end = next_starts[0] if next_starts else len(full_text)
        block = full_text[start:end].strip()
        if len(block) > 80:
            blocks.append(block)

    return blocks

def _context_passage_for_trigger(text: str, trigger: str, max_chars: int = 600) -> str:
    """
    Returns the sentence before, the sentence containing trigger, and the sentence after.
    Plain text — bolding is applied at render time in render_finding_card.
    Falls back to a character window if sentence split fails.
    """
    sentences = _split_into_sentences(text)
    trigger_lower = trigger.lower()

    target_idx = None
    for i, sent in enumerate(sentences):
        if trigger_lower in sent.lower():
            target_idx = i
            break

    if target_idx is None:
        m = re.search(
            rf'.{{0,200}}\b{re.escape(trigger)}\b.{{0,200}}',
            text, re.IGNORECASE | re.DOTALL
        )
        if m:
            return re.sub(r'\s+', ' ', m.group(0)).strip()[:max_chars]
        return trigger

    start_idx = max(0, target_idx - 1)
    end_idx   = min(len(sentences), target_idx + 2)
    passage   = ' '.join(sentences[start_idx:end_idx])
    return re.sub(r'\s+', ' ', passage).strip()[:max_chars]
def _sentence_for_trigger(text: str, trigger: str, max_chars: int = 300) -> str:
    """
    Returns the complete sentence from text that contains trigger.
    Cleans whitespace. Falls back to character window if sentence split fails.
    """
    sentences = _split_into_sentences(text)
    trigger_lower = trigger.lower()
    for sent in sentences:
        if trigger_lower in sent.lower():
            clean = re.sub(r'\s+', ' ', sent).strip()
            return clean[:max_chars]
    m = re.search(
        rf'.{{0,150}}\b{re.escape(trigger)}\b.{{0,150}}',
        text, re.IGNORECASE | re.DOTALL
    )
    if m:
        return re.sub(r'\s+', ' ', m.group(0)).strip()[:max_chars]
    return trigger

def _split_into_sentences(text: str) -> list:
    raw = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    return [s.strip() for s in raw if s.strip()]


def _extract_key_terms(phrase: str) -> list:
    words = re.findall(r"[a-zA-Z']+", phrase.lower())
    return [w for w in words if len(w) > 3 and w not in _REC_STOPWORDS]


def _will_sentences_from_section_f(section_f_text: str) -> list:
    sentences = _split_into_sentences(section_f_text)
    will_re = re.compile(r'\bwill\b', re.IGNORECASE)
    return [s for s in sentences if will_re.search(s)]


def _score_match(key_terms: list, candidate_sentence: str) -> tuple:
    lower = candidate_sentence.lower()
    matched = [t for t in key_terms if t in lower]
    return len(matched), matched
    
def analyse_section_f(f_blocks):
    findings = []

    if not f_blocks:
        findings.append({
            "tier": "red",
            "title": "Section F not located",
            "extract": "No Section F provision content was identified in this document.",
            "commentary": (
                "Section F is a lawful requirement under the Children and Families Act 2014. "
                "Every EHCP must specify the educational provision in Section F. "
                "If this is a draft document, check whether the section has been completed. "
                "If this is a final EHCP, raise the absence at the earliest opportunity."
            ),
            "delivery_log_required": True,
        })
        return findings

    combined_f = "\n\n".join(f_blocks)

    freq_pattern = re.compile(r'\b(\d+)\s*(session|hour|minute|time|per week|per term|weekly|daily)\b', re.IGNORECASE)
    has_frequency = bool(freq_pattern.search(combined_f))
    if not has_frequency:
        findings.append({
            "tier": "red",
            "title": "Frequency of provision not specified",
            "extract": _get_short_extract(combined_f, 0, 300),
            "commentary": (
                "Section F must state how often each provision is delivered — "
                "for example, three sessions per week. Without a stated frequency, "
                "the provision cannot be lawfully enforced under the Children and Families Act 2014. "
                "Raise at annual review and request specific numeric frequency for each provision."
            ),
            "delivery_log_required": True,
        })

    dur_pattern = re.compile(r'\b(\d+)\s*(minute|hour|min)\b', re.IGNORECASE)
    has_duration = bool(dur_pattern.search(combined_f))
    if not has_duration:
        findings.append({
            "tier": "red",
            "title": "Duration of sessions not specified",
            "extract": _get_short_extract(combined_f, 0, 300),
            "commentary": (
                "Section F must state the length of each session in minutes or hours. "
                "Without a stated duration, provision is unquantifiable and unenforceable. "
                "Request an amendment to specify session length for every provision."
            ),
            "delivery_log_required": True,
        })

    role_pattern = re.compile(
        r'\b(SENCO|teaching assistant|TA|therapist|specialist|speech|occupational|'
        r'psychologist|learning support|LSA|teacher|practitioner|coordinator)\b',
        re.IGNORECASE
    )
    has_role = bool(role_pattern.search(combined_f))
    if not has_role:
        findings.append({
            "tier": "red",
            "title": "Deliverer role not specified",
            "extract": _get_short_extract(combined_f, 0, 300),
            "commentary": (
                "Section F must identify who delivers each provision — the role title "
                "and relevant qualification or training level. This is a lawful requirement. "
                "Without a named role, accountability for delivery cannot be established."
            ),
            "delivery_log_required": True,
        })

    # ── Check Group 1: Modal language — commitment strength ───────────────────
    # "will benefit from", "should receive", "may be provided" etc.
    # These replace a lawful commitment with a conditional or aspirational statement.
    found_modal = []
    for word in VAGUE_MODAL:
        if re.search(rf'\b{re.escape(word)}\b', combined_f, re.IGNORECASE):
           ctx = _sentence_for_trigger(combined_f, word)
           found_modal.append((word, ctx))

    if found_modal:
        examples = "; ".join(f'"{w}"' for w, _ in found_modal[:4])
        findings.append({
            "tier": "red",
            "title": "Section F uses conditional language instead of a lawful commitment",
            "extract": _context_passage_for_trigger(combined_f, found_modal[0][0]),
            "trigger_word": found_modal[0][0],
            "commentary": (
                "Section F must state what provision 'will' be delivered — not what the child "
                "'should receive', 'would benefit from', or 'may be provided with'. "
                "A provision written as a recommendation or a possibility is not a legal "
                "commitment and cannot be enforced via Judicial Review. "
                "Under the Children and Families Act 2014, the LA has an absolute duty to "
                "secure whatever is specified in Section F. Only 'will receive' creates that duty. "
                "Each instance must be challenged and rewritten as a specific, unambiguous "
                "commitment before the plan is finalised."
            ),
            "delivery_log_required": True,
        })

    # ── Check Group 2: Vague qualifiers — specificity ────────────────────────
    # "access to", "as appropriate", "regular", "up to" etc.
    # SOS!SEN explicit list — each one makes provision impossible to audit or enforce.
    found_qualifier = []
    for word in VAGUE_QUALIFIER:
        if re.search(rf'\b{re.escape(word)}\b', combined_f, re.IGNORECASE):
           ctx = _sentence_for_trigger(combined_f, word)
           found_qualifier.append((word, ctx))

    if found_qualifier:
        examples = "; ".join(f'"{w}"' for w, _ in found_qualifier[:4])
        findings.append({
            "tier": "red",
            "title": "Section F contains vague qualifier language that removes enforceability",
            "extract": _context_passage_for_trigger(combined_f, found_qualifier[0][0]),
            "trigger_word": found_qualifier[0][0],
            "commentary": (
                "The SEN Code of Practice (paragraph 9.69) requires Section F to be "
                "detailed and specific — and normally quantified. Terms like 'access to', "
                "'as appropriate', 'where necessary', 'regular', and 'up to X hours' "
                "make it impossible to establish whether provision has been delivered. "
                "'Regular' must become 'three sessions per week'. "
                "'Access to' must become 'will receive'. "
                "'Up to X hours' must become a specific commitment, not an upper limit. "
                "Every vague qualifier in Section F is a gap in enforceability and "
                "should be challenged at or before the annual review."
            ),
            "delivery_log_required": True,
        })

    log_pattern = re.compile(r'\b(delivery log|record of delivery|session record|log)\b', re.IGNORECASE)
    has_log_ref = bool(log_pattern.search(combined_f))
    if not has_log_ref:
        findings.append({
            "tier": "amber",
            "title": "No delivery log mechanism referenced",
            "extract": "",
            "commentary": (
                "Best practice requires that provision delivery is logged contemporaneously. "
                "A delivery log is the evidence base for the Do stage of the APDR cycle. "
                "Without a delivery log, the Review stage cannot be conducted accurately. "
                "Ask the school what logging system is in place and request access to records."
            ),
            "delivery_log_required": True,
        })

    universal_pattern = re.compile(
        r'\b(quality first teaching|ordinarily available provision|'
        r'universal provision|whole school approach|available to all pupils)\b',
        re.IGNORECASE
    )
    m = universal_pattern.search(combined_f)
    if m:
        ctx = combined_f[max(0, m.start()-80):m.end()+80].strip()
        disclaimer = re.compile(
            r'provision set out here is in addition to that which all pupils',
            re.IGNORECASE
        )
        if not disclaimer.search(ctx):
            findings.append({
                "tier": "red",
                "title": "Universal provision substituted for specific provision",
                "extract": ctx[:300],
                "commentary": (
                    "Section F must contain provision that is specific to this child, "
                    "above and beyond what is available to all pupils. Universal or "
                    "'quality first' provision does not meet the lawful requirement for "
                    "specified provision in an EHCP. This is the most serious Section F failure. "
                    "Challenge and request replacement with individually specified provision."
                ),
                "delivery_log_required": True,
            })

    launder_pattern = re.compile(
        r'\b(it is recommended|it would be beneficial|it is suggested|'
        r'consideration should be given|it is hoped)\b',
        re.IGNORECASE
    )
    m = launder_pattern.search(combined_f)
    if m:
        ctx = combined_f[max(0, m.start()-80):m.end()+80].strip()
        findings.append({
            "tier": "red",
            "title": "Recommendation language detected — not converted to lawful commitment",
            "extract": ctx[:300],
            "commentary": (
                "Assessment language has been copied into Section F without being converted "
                "into a specified lawful commitment. Section F must contain commitments, "
                "not recommendations. 'It is recommended' has no enforceability. "
                "Each recommendation must be rewritten as a specified, quantified commitment."
            ),
            "delivery_log_required": False,
        })

    # Tighter dilution clause — only fires on conditional constructions,
    # not on standalone words like "resources" or "staffing arrangements".
    # Dilution clause — tighter pattern, conditional constructions only.
    # Extended with Provision Substitution terms from FRED_PROCESS_PATTERNS.
    dilution_pattern = re.compile(
        r'\b(?:'
        r'subject to (?:available )?(?:resources|staffing|funding|budget|capacity)|'
        r'dependent on (?:available )?(?:resources|staffing|availability|funding)|'
        r'contingent on (?:available )?(?:resources|staffing|budget|funding)|'
        r'where resources (?:allow|permit)|'
        r'where staffing (?:allows|permits)|'
        r'if staffing (?:allows|permits)|'
        r'where capacity (?:allows|permits)|'
        r'subject to funding|'
        r'resources permitting|'
        r'if (?:resources|budget) (?:allow|permits?)|'
        r'within available resources|'
        r'delivered operational capacity permitting|'
        r'resource-dependent intervention|'
        r'capacity-linked support|'
        r'resource-matched intervention|'
        r'subject to specialist availability|'
        r'subject to recruitment timelines|'
        r'delivered within existing frameworks|'
        r'delivered within standard ratios|'
        r'delivered via general school environment|'
        r'operational adjustments applied|'
        r'pro-rata provision delivery|'
        r'resource allocation panel decision|'
        r'delivered by available pastoral staff|'
        r'integrated within standard delivery|'
        r'resource-limited intervention'
        r')\b',
        re.IGNORECASE
    )
    m = dilution_pattern.search(combined_f)
    if m:
        ctx = combined_f[max(0, m.start()-80):m.end()+80].strip()
        findings.append({
            "tier": "red",
            "title": "Dilution clause detected — provision made conditional on resources or staffing",
            "extract": ctx[:300],
            "commentary": (
                "Section F provision must not be conditional on school resources, staffing, "
                "or budget. The duty to deliver specified provision under the Children and "
                "Families Act 2014 is absolute — it does not depend on what the school can "
                "afford or what staff are available on a given day. "
                "If additional funding is required to deliver what is in Section F, "
                "that obligation falls on the LA under s42 CFA 2014, not the school. "
                "Any dilution clause must be identified, challenged, and removed before "
                "the plan is finalised."
            ),
            "delivery_log_required": False,
        })
    m = dilution_pattern.search(combined_f)
    if m:
        ctx = combined_f[max(0, m.start()-80):m.end()+80].strip()
        findings.append({
            "tier": "red",
            "title": "Dilution clause detected — provision made conditional on resources or staffing",
            "extract": ctx[:300],
            "commentary": (
                "Section F provision must not be conditional on school resources, staffing, "
                "or budget. The duty to deliver specified provision under the Children and "
                "Families Act 2014 is absolute — it does not depend on what the school can "
                "afford or what staff are available on a given day. "
                "If additional funding is required to deliver what is in Section F, "
                "that obligation falls on the LA under s42 CFA 2014, not the school. "
                "Any dilution clause must be identified, challenged, and removed before "
                "the plan is finalised."
            ),
            "delivery_log_required": False,
        })

    n_red = sum(1 for f in findings if f["tier"] == "red")
    if n_red == 0:
        findings.append({
            "tier": "green",
            "title": "Section F provision appears lawfully specified",
            "extract": "",
            "commentary": (
                "Section F contains stated frequency, duration, and deliverer role. "
                "No prohibited language was detected. Use these compliant entries as "
                "benchmarks when challenging any non-compliant provision in this or "
                "future EHCPs."
            ),
            "delivery_log_required": False,
        })

    return findings
def analyse_report_recommendations(report_text: str, section_f_text: str,
                                   report_label: str = "Professional report") -> list:
    if not report_text or not section_f_text:
        return []

    findings = []
    sentences = _split_into_sentences(report_text)
    will_sentences = _will_sentences_from_section_f(section_f_text)

    trigger_pattern = re.compile(
        r'(' + '|'.join(re.escape(t) for t in RECOMMENDATION_TRIGGERS) + r')',
        re.IGNORECASE
    )

    seen = set()

    for sentence in sentences:
        m = trigger_pattern.search(sentence)
        if not m:
            continue

        key = sentence.strip().lower()
        if key in seen:
            continue
        seen.add(key)

        trigger_found = m.group(1)
        substance = sentence[m.end():].strip()
        key_terms = _extract_key_terms(substance)

        if len(key_terms) < 2:
            continue

        best_score = 0
        best_sentence = None
        best_matched = []

        for will_s in will_sentences:
            score, matched = _score_match(key_terms, will_s)
            if score > best_score:
                best_score = score
                best_sentence = will_s
                best_matched = matched

        if best_score >= _REC_MATCH_THRESHOLD:
            findings.append({
                "type": "GREEN",
                "report_label": report_label,
                "recommendation": sentence.strip(),
                "trigger": trigger_found,
                "key_terms": key_terms,
                "matched_terms": best_matched,
                "section_f_match": best_sentence,
                "finding": (
                    "A recommendation from this report appears to have been carried "
                    "through into Section F as a commitment. That is what good looks "
                    "like — a professional identifies a need, and the EHCP converts it "
                    "into something the school is legally required to deliver. "
                    "Check the Section F wording carefully. It must state what will be "
                    "delivered, how often, for how long, and by whom. If it uses vague "
                    "language — 'as appropriate', 'access to', 'regularly' — that will "
                    "appear as a separate red finding."
                ),
                "citation": "",
            })
        else:
            findings.append({
                "type": "RED",
                "report_label": report_label,
                "recommendation": sentence.strip(),
                "trigger": trigger_found,
                "key_terms": key_terms,
                "matched_terms": [],
                "section_f_match": None,
                "finding": (
                    "A professional who assessed your child recommended this provision. "
                    "It does not appear in Section F as a legal commitment — which means "
                    "the school and the LA have no duty to deliver it. "
                    "A recommendation that stays in an assessment report but never makes "
                    "it into Section F is not enforceable. "
                    "At your next annual review, ask directly: why has this not been "
                    "written into the plan? If there is no satisfactory answer, request "
                    "that it is added before the review closes."
                ),
                "citation": (
                    "Children and Families Act 2014 s42 — the duty to deliver provision "
                    "is absolute, but only for provision that is specified in Section F. "
                    "A recommendation that has not been converted to a 'will' commitment "
                    "in Section F is not a legal entitlement and the LA can decline to act on it."
                ),
            })

    return findings
    
def analyse_section_e(e_blocks):
    findings = []

    if not e_blocks:
        findings.append({
            "tier": "amber",
            "title": "Section E outcomes not located",
            "extract": "",
            "commentary": (
                "Section E should contain SMART outcomes — specific, measurable, "
                "achievable, realistic, and time-limited. Each outcome should describe "
                "what the child will be able to do as a result of the provision in "
                "Section F, not what they will receive. If outcomes are absent from "
                "the EHCP, raise this at annual review and request that Section E is "
                "completed before the plan is finalised."
            ),
            "delivery_log_required": False,
        })
        return findings

    combined_e = "\n\n".join(e_blocks)

    # ── Check 1: Timeframe ────────────────────────────────────────────────────
    # EHCP-relevant timeframes only — not APDR targets or school-based plans.
    # Recognises: Key Stage references, Year group references, annual review
    # references, age references, transition references, and dated timeframes.
    time_pattern = re.compile(
        r'\b('
        # Dated timeframes — "by July 2025", "by end of 2026"
        r'by \w+ 20\d\d|'
        r'within \d+ (?:month|week|term|year)s?|'
        # Key Stage references
        r'by (?:the )?end of (?:the )?(?:key stage|ks)\s*[1-4]|'
        r'before (?:the )?end of (?:the )?(?:key stage|ks)\s*[1-4]|'
        r'before (?:the )?end of (?:key stage|ks)\s*[1-4]|'
        # Year group references
        r'by (?:the )?end of year \d+|'
        r'by year \d+|'
        # Annual review references
        r'by (?:the )?next annual review|'
        r'at (?:the )?next annual review|'
        r'by (?:the )?annual review|'
        # Age references
        r'by (?:the )?age of \d+|'
        r'by age \d+|'
        # Transition references
        r'before (?:transition|transferring|moving) to|'
        r'by (?:the )?time (?:of|they reach)|'
        # Phase references
        r'by (?:the )?end of (?:primary|secondary|sixth form|college)|'
        r'before (?:leaving|starting) (?:primary|secondary|sixth form|college)'
        r')\b',
        re.IGNORECASE
    )
    has_timeframe = bool(time_pattern.search(combined_e))

    if not has_timeframe:
        findings.append({
            "tier": "red",
            "title": "Section E outcomes lack a measurable timeframe",
            "extract": _get_short_extract(combined_e, 0, 300),
            "commentary": (
                "Every outcome in Section E must include a timeframe — when the child "
                "is expected to achieve it. The SEN Code of Practice requires outcomes "
                "to be time-limited. Acceptable timeframes in an EHCP include: "
                "'by the end of Key Stage 2', 'by the next annual review', "
                "'by July 2026', 'by age 11', or 'before transition to secondary'. "
                "Without a timeframe, the annual review cannot establish whether an "
                "outcome has been achieved. This must be raised and corrected before "
                "the plan is finalised."
            ),
            "delivery_log_required": False,
        })

# ── Check 2: Achievement vs activity language ─────────────────────────────
    # CoP para 9.69: outcomes describe what the child will be ABLE TO DO,
    # not what they will RECEIVE. "Will receive" is provision language —
    # it belongs in Section F, not Section E.

    achievement_pattern = re.compile(
        r'\b('
        r'will be able to|'
        r'will independently|'
        r'will demonstrate|'
        r'will achieve|'
        r'will show|'
        r'will communicate|'
        r'will manage|'
        r'will use|'
        r'will apply|'
        r'will self.regulate|'
        r'will initiate|'
        r'will sustain|'
        r'will access independently|'
        r'will produce|'
        r'will read|'
        r'will write|'
        r'will participate'
        r')\b',
        re.IGNORECASE
    )

    activity_pattern = re.compile(
        r'\b('
        r'will receive|'
        r'will attend|'
        r'will be supported|'
        r'will be provided|'
        r'will have access|'
        r'will be given|'
        r'will be included|'
        r'will be helped|'
        r'will be encouraged|'
        r'will be offered|'
        r'will be placed|'
        r'will be monitored|'
        r'will be reviewed|'
        r'will take part in|'
        r'will participate in sessions|'
        r'will be exposed to'
        r')\b',
        re.IGNORECASE
    )

    # ── Check 3: Vague aspiration language ───────────────────────────────────
    # "Will have developed", "will have improved", "will have made progress" —
    # directional language that describes a state rather than a capability.
    # Cannot be measured at annual review. Protects the school from accountability
    # while appearing child-centred. Always benefits the institution, never the child.
    vague_aspiration_pattern = re.compile(
        r'\b('
        r'will have developed|'
        r'will have improved|'
        r'will have increased|'
        r'will have made progress|'
        r'will have gained|'
        r'will have grown|'
        r'will have built|'
        r'will have enhanced|'
        r'will have extended|'
        r'will have progressed|'
        r'will have strengthened|'
        r'will have consolidated|'
        r'will have broadened|'
        r'will have begun to|'
        r'will have started to|'
        r'will have continued to|'
        r'will have worked towards|'
        r'will have explored|'
        r'will have experienced|'
        r'will have been introduced to'
        r')\b',
        re.IGNORECASE
    )

    has_achievement  = bool(achievement_pattern.search(combined_e))
    has_activity     = bool(activity_pattern.search(combined_e))
    has_aspiration   = bool(vague_aspiration_pattern.search(combined_e))

    # Activity language finding
    if has_activity and not has_achievement:
        m = activity_pattern.search(combined_e)
        ctx = combined_e[max(0, m.start()-80):m.end()+120].strip() if m else ""
        findings.append({
            "tier": "red",
            "title": "Section E outcomes written as activities, not achievements",
            "extract": ctx[:300],
            "commentary": (
                "Section E outcomes must describe what the child will be able to do — "
                "not what they will receive or attend. This is a requirement of the "
                "SEN Code of Practice (paragraph 9.69). "
                "'Freddie will receive speech and language therapy' is provision — "
                "it belongs in Section F. "
                "'Freddie will be able to use two-word combinations to make requests' "
                "is an outcome — it belongs in Section E. "
                "Outcomes written as activities cannot be measured at annual review "
                "and make the EHCP significantly harder to enforce. "
                "Every outcome must describe an achievement, not a service."
            ),
            "delivery_log_required": False,
        })
    elif has_activity and has_achievement:
        m = activity_pattern.search(combined_e)
        ctx = combined_e[max(0, m.start()-80):m.end()+120].strip() if m else ""
        findings.append({
            "tier": "amber",
            "title": "Section E contains a mix of achievement and activity language",
            "extract": ctx[:300],
            "commentary": (
                "Some outcomes in Section E are written correctly as achievements — "
                "describing what the child will be able to do. Others are written as "
                "activities — describing what the child will receive or attend. "
                "Activity-language outcomes cannot be measured at annual review and "
                "should be rewritten. At annual review, ask for each activity-language "
                "outcome to be replaced with an achievement: not 'will attend a social "
                "skills group' but 'will be able to initiate a conversation with a peer "
                "in a structured setting'."
            ),
            "delivery_log_required": False,
        })

    # Vague aspiration language finding — separate RED, separate explanation
    if has_aspiration:
        m = vague_aspiration_pattern.search(combined_e)
        ctx = combined_e[max(0, m.start()-80):m.end()+120].strip() if m else ""
        findings.append({
            "tier": "red",
            "title": "Section E outcomes use vague aspiration language — unmeasurable at annual review",
            "extract": ctx[:300],
            "commentary": (
                "Outcomes like 'will have developed', 'will have improved', or "
                "'will have made progress' describe a direction of travel, not an "
                "achievable, measurable endpoint. They cannot be objectively assessed "
                "at annual review — the school can always claim some development has "
                "occurred without specifying what was expected. "
                "This language benefits the institution, not the child. "
                "A measurable outcome names a specific capability, a context, and "
                "a standard: 'will be able to spend 20 minutes in a group of three "
                "peers and initiate conversation on two occasions without adult "
                "prompting'. That outcome can be assessed. 'Will have developed "
                "his social understanding' cannot. "
                "Every vague aspiration outcome must be rewritten at annual review "
                "with a specific capability, a measurable standard, and a timeframe."
            ),
            "delivery_log_required": False,
        })

    # ── Check 3: SMART language — measurability ───────────────────────────────
    # Checks whether outcomes contain any measurable indicator.
    # Not checking for baselines — those belong in assessment reports, not the EHCP.
    measurable_pattern = re.compile(
        r'\b('
        r'\d+|'                          # any number
        r'independently|'
        r'without support|'
        r'without prompting|'
        r'with minimal|'
        r'on \d+ out of|'
        r'in \d+ out of|'
        r'four out of five|'
        r'three out of four|'
        r'consistently|'
        r'across all|'
        r'in all settings|'
        r'measured by|'
        r'evidenced by|'
        r'assessed by'
        r')\b',
        re.IGNORECASE
    )
    has_measurable = bool(measurable_pattern.search(combined_e))

    if not has_measurable:
        findings.append({
            "tier": "amber",
            "title": "Section E outcomes may lack a measurable success criterion",
            "extract": _get_short_extract(combined_e, 0, 300),
            "commentary": (
                "The SEN Code of Practice requires outcomes to be measurable — "
                "there should be a way to establish at annual review whether the "
                "outcome has been achieved. Measurable outcomes include specific "
                "numbers ('four out of five occasions'), independence indicators "
                "('without adult prompting'), or setting criteria ('across all "
                "lessons'). A vague outcome like 'Freddie will improve his "
                "communication' cannot be objectively reviewed. At annual review, "
                "ask for each outcome to include a specific success criterion."
            ),
            "delivery_log_required": False,
        })

    # ── Green — outcomes appear correctly structured ───────────────────────────
    n_red = sum(1 for f in findings if f["tier"] == "red")
    if n_red == 0:
        findings.append({
            "tier": "green",
            "title": "Section E outcomes appear correctly structured",
            "extract": "",
            "commentary": (
                "Section E contains timeframe indicators and achievement language. "
                "Outcomes appear to describe what the child will be able to do "
                "rather than what they will receive. Use these as benchmarks when "
                "challenging any outcomes that do not meet this standard in future "
                "reviews or amended plans."
            ),
            "delivery_log_required": False,
        })

    return findings

def _build_provision_regex(category: str) -> re.Pattern:
    """
    Builds a compiled regex from PROVISION_LIBRARY terms for the given category.
    Used in analyse_section_b() as a second-stage precision check.
    Returns a pattern that matches any known provision term for that category.
    Falls back to an never-matching pattern if category is not found.
    """
    terms = PROVISION_LIBRARY.get(category, [])
    if not terms:
        return re.compile(r"(?!)")  # never matches
    escaped = [re.escape(term) for term in terms]
    pattern_str = r'\b(?:' + "|".join(escaped) + r')\b'
    return re.compile(pattern_str, re.IGNORECASE)

def analyse_section_b(b_blocks, f_blocks):
    findings = []
    if not b_blocks:
        return findings

    combined_b = "\n\n".join(b_blocks)
    combined_f = "\n\n".join(f_blocks) if f_blocks else ""

    need_areas = [
        {
            "name": "Communication and interaction",
            "library_key": "Communication and Interaction",
            "b_keywords": r'\b(communication|speech|language|SALT|interaction|social communication|autism|ASD|pragmatic)\b',
            "f_keywords": r'\b(speech|language|communication|SALT|social skills|interaction)\b',
            "what_good_looks_like": (
                "Section F should specify: the frequency of speech and language therapy or "
                "targeted communication sessions (e.g. two sessions per week), the duration "
                "of each session in minutes, the role and qualification of the deliverer "
                "(e.g. Speech and Language Therapist or trained TA working to a SALT programme), "
                "and the specific strategies to be used. Generalised statements such as "
                "'communication support will be provided' are not sufficient."
            ),
        },
        {
            "name": "Cognition and learning",
            "library_key": "Cognition and Learning",
            "b_keywords": r'\b(cognition|learning|literacy|numeracy|reading|writing|dyslexia|processing|memory|attention)\b',
            "f_keywords": r'\b(literacy|numeracy|reading|writing|learning support|intervention|programme)\b',
            "what_good_looks_like": (
                "Section F should specify: named literacy or learning interventions (not generic "
                "'learning support'), frequency and duration of sessions, the role of the "
                "deliverer and their relevant training, and measurable targets that link "
                "directly to Section E outcomes. 'Access to' a resource is not provision."
            ),
        },
        {
            "name": "Social, emotional and mental health",
            "library_key": "SEMH",
            "b_keywords": r'\b(SEMH|emotional|mental health|anxiety|behaviour|wellbeing|regulation|self.esteem|ADHD|attachment)\b',
            "f_keywords": r'\b(emotional|SEMH|regulation|therapeutic|counselling|pastoral|wellbeing|anxiety)\b',
            "what_good_looks_like": (
                "Section F should specify: named therapeutic or emotional support provision "
                "(e.g. weekly one-to-one sessions with a trained SEMH TA or school counsellor), "
                "the frequency and duration, the deliverer's role and relevant training, "
                "and specific strategies referenced from assessment. 'Pastoral support' "
                "without specification does not meet the lawful standard."
            ),
        },
        {
            "name": "Sensory and physical needs",
            "library_key": "Sensory and Physical",
            "b_keywords": r'\b(sensory|physical|motor|OT|occupational therapy|fine motor|gross motor|proprioception|vestibular|sensory processing)\b',
            "f_keywords": r'\b(OT|occupational|sensory|motor|physical|sensory diet|movement break)\b',
            "what_good_looks_like": (
                "Section F should specify: named occupational therapy or sensory integration "
                "provision, frequency and duration of sessions, the role of the deliverer "
                "(e.g. Occupational Therapist or TA trained to deliver OT programme), "
                "and the specific programme or strategies. Environmental adjustments should "
                "be listed specifically, not described as generalised 'reasonable adjustments'."
            ),
        },
    ]

    for area in need_areas:
        b_match = re.search(area["b_keywords"], combined_b, re.IGNORECASE)

        # Two-stage F check:
        # Stage 1 — broad keyword match (topic present)
        # Stage 2 — PROVISION_LIBRARY term match (named provision present)
        f_broad  = re.search(area["f_keywords"], combined_f, re.IGNORECASE) if combined_f else None
        f_library = _build_provision_regex(area["library_key"]).search(combined_f) if combined_f else None
        f_match  = f_broad or f_library

        if b_match and not f_match:
            ctx = combined_b[max(0, b_match.start()-60):b_match.end()+120].strip()
            findings.append({
                "tier": "red",
                "title": f"Need identified in Section B with no provision in Section F — {area['name']}",
                "extract": ctx[:280],
                "commentary": (
                    f"Section B identifies {area['name'].lower()} needs but Section F contains "
                    f"no corresponding provision. Under the Children and Families Act 2014, "
                    f"every need identified in Section B must be met by provision in Section F. "
                    f"This gap must be addressed at annual review.\n\n"
                    f"What good looks like: {area['what_good_looks_like']}"
                ),
                "delivery_log_required": True,
            })

        elif b_match and f_match:
            f_ctx = combined_f[max(0, f_match.start()-60):f_match.end()+120].strip()
            vague = re.search(
                r'\b(support will be provided|will be supported|will have access|'
                r'as appropriate|where necessary|as needed)\b',
                f_ctx, re.IGNORECASE
            )
            if vague:
                findings.append({
                    "tier": "amber",
                    "title": f"Section F provision for {area['name']} may lack specificity",
                    "extract": f_ctx[:280],
                    "commentary": (
                        f"Section B identifies {area['name'].lower()} needs and Section F "
                        f"references corresponding provision, but the language used may be "
                        f"insufficiently specific to be enforceable.\n\n"
                        f"What good looks like: {area['what_good_looks_like']}"
                    ),
                    "delivery_log_required": False,
                })

    health_pattern = re.compile(r'\b(health|medical|therapy|therapist|OT|SALT|physiotherapy)\b', re.IGNORECASE)
    if health_pattern.search(combined_b) and not health_pattern.search(combined_f):
        findings.append({
            "tier": "amber",
            "title": "Health needs in Section B — check Sections C and G for corresponding provision",
            "extract": "",
            "commentary": (
                "Section B describes health-related needs. Sections C and G should contain "
                "corresponding health provision and outcomes. If these sections are empty or "
                "absent, the needs identified in Section B are not being addressed through "
                "the full EHCP framework. Raise at annual review."
            ),
            "delivery_log_required": False,
        })

    return findings


def _get_short_extract(text, start, length):
    extract = text[start:start+length].strip()
    extract = re.sub(r'\n{3,}', '\n\n', extract)
    extract = re.sub(r' {2,}', ' ', extract)
    return extract

def extract_la_name(full_text: str) -> str:
    """
    Best-effort extraction of the local authority name from EHCP text.
    Looks in the first 3000 characters where the issuing body typically appears.
    Returns a string if found, empty string if not.
    """
    search_window = full_text[:3000]

    # Pattern 1: "Local Authority: Warwickshire" or "Local Authority — Warwickshire"
    m = re.search(
        r'local authority\s*[:\-–]\s*([A-Z][a-zA-Z\s]{2,40}?)(?:\n|,|\.)',
        search_window, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()

    # Pattern 2: "Warwickshire County Council" or "Birmingham City Council"
    m = re.search(
        r'\b([A-Z][a-zA-Z\s]{2,30}(?:County|City|Borough|District|Metropolitan)\s+Council)\b',
        search_window
    )
    if m:
        return m.group(1).strip()

    # Pattern 3: "issued by Warwickshire" or "prepared by Warwickshire"
    m = re.search(
        r'(?:issued|prepared|produced)\s+by\s+([A-Z][a-zA-Z\s]{2,30}?)(?:\n|,|\.)',
        search_window, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()

    return ""

def extract_child_name(full_text: str) -> str:
    name_label_pattern = re.compile(
        r"(?:child'?s?\s+name|name\s+of\s+child|full\s+name|pupil'?s?\s+name|student'?s?\s+name)"
        r"\s*[:\-]?\s*([A-Z][a-z]{1,20}(?:\s+[A-Z][a-z]{1,20}){0,3})",
        re.IGNORECASE
    )
    m = name_label_pattern.search(full_text)
    if m:
        first_name = m.group(1).strip().split()[0].lower()
        if len(first_name) > 2:
            return first_name
    dob_marker = re.compile(
        r'\b(date\s+of\s+birth|dob|date\s+issued)\b',
        re.IGNORECASE
    )
    dob_m = dob_marker.search(full_text)
    if dob_m:
        start = max(0, dob_m.start() - 300)
        end   = min(len(full_text), dob_m.end() + 300)
        window = full_text[start:end]
        name_in_window = re.search(
            r'\b([A-Z][a-z]{2,20})\s+([A-Z][a-z]{2,20})\b',
            window
        )
        if name_in_window:
            candidate = name_in_window.group(1).lower()
            excluded = {"date", "name", "birth", "section", "plan", "school",
                        "local", "authority", "education", "health", "care"}
            if candidate not in excluded:
                return candidate
    return ""
                
def build_ehcp_commitments_summary(f_blocks: list) -> list:
    """
    Extracts plain-English commitment sentences from Section F blocks.
    Returns up to 20 cleaned, deduplicated sentences containing 'will'.
    Filters out administrative, structural, and deliverer-column sentences.
    Used to display 'Your child's EHCP commits to:' at the top of the report.
    """
    if not f_blocks:
        return []

    ADMIN_PATTERNS = re.compile(
        r'\b(who will provide|provision set out|all pupils|in addition to|'
        r'communication friendly|quality first|ordinarily available|'
        r'deliverer|section f|section e|section b|annual review|'
        r'children and families act|send code|local authority|'
        r'this plan|this ehcp|the ehcp|named school|placement|'
        r'will be discharged|will benefit his|will benefit her)\b',
        re.IGNORECASE
    )

    MIN_WORDS = 6
    MAX_WORDS = 50

    seen = set()
    results = []

    for block in f_blocks:
        sentences = _split_into_sentences(block)
        for sent in sentences:
            # Must contain 'will'
            if not re.search(r'\bwill\b', sent, re.IGNORECASE):
                continue
            # Filter admin/structural sentences
            if ADMIN_PATTERNS.search(sent):
                continue
            # Clean whitespace
            clean = re.sub(r'\s+', ' ', sent).strip()
            # Length gate
            word_count = len(clean.split())
            if word_count < MIN_WORDS or word_count > MAX_WORDS:
                continue
            # Deduplicate on normalised lowercase
            key = clean.lower()
            if key in seen:
                continue
            seen.add(key)
            results.append(clean)
            if len(results) >= 20:
                return results

    return results
def write_analysis_to_supabase(findings: list, meta: dict, la_name: str = None):
    if not SUPABASE_AVAILABLE or not supabase:
        return False
    st.markdown("<script>window.scrollTo(0, 0);</script>", unsafe_allow_html=True)
    user = st.session_state.get("user")
    if not user:
        return False
    try:
        import json as _json
        user_id = user.id
        supabase.auth.set_session(st.session_state["session"].access_token, st.session_state["session"].refresh_token)
        today_label = datetime.datetime.now().strftime("%d %b %Y")
        document_label = f"EHCP uploaded {today_label}"
        supabase.table("document_logs").insert({
            "user_id": str(user_id),
            "document_label": document_label,
            "la_name": st.session_state.get("la_name", None) or None,
        }).execute()
        red_count = sum(1 for f in findings if f.get("tier") == "red")
        vague_count = sum(1 for f in findings if "vague" in f.get("title", "").lower())
        commitments_count = len(
            build_ehcp_commitments_summary(
                find_section_blocks(st.session_state.get("full_text", ""), "F")
            )
        )
        supabase.table("analysis_findings").insert({
            "user_id": str(user_id),
            "section_f_blocks_found": meta.get("f_blocks_found", 0),
            "vague_terms_count": vague_count,
            "red_findings_count": red_count,
            "commitments_count": commitments_count,
            "raw_findings_json": _json.dumps(findings, default=str),
        }).execute()
        return True
    except Exception as e:
        st.warning(f"Supabase write failed: {e}")
        return False
    
def run_full_analysis(full_text):
    f_blocks = find_section_blocks(full_text, "F")
    e_blocks = find_section_blocks(full_text, "E")
    b_blocks = find_section_blocks(full_text, "B")

    child_name = extract_child_name(full_text)
    la_name = extract_la_name(full_text)
    st.session_state["la_name"] = la_name
    if child_name:
        _REC_STOPWORDS.add(child_name)
        st.session_state["child_name"] = child_name
    else:
        st.session_state["child_name"] = ""

    f_findings = analyse_section_f(f_blocks)
    e_findings = analyse_section_e(e_blocks)
    b_findings = analyse_section_b(b_blocks, f_blocks)

    all_findings = f_findings + e_findings + b_findings

    order = {"red": 0, "amber": 1, "green": 2}
    all_findings.sort(key=lambda x: order.get(x["tier"], 3))

    return all_findings, {
        "f_blocks_found": len(f_blocks),
        "e_blocks_found": len(e_blocks),
        "b_blocks_found": len(b_blocks),
    }


# ── REPORT GENERATORS ─────────────────────────────────────────────────────────

def generate_word_report(
    findings,
    child_name="your child",
    situation="",
    doc_type="EHCP",
    commitments=None
):
    doc = Document()

    title = doc.add_heading("FRED Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Families' Rights and Entitlements Directory")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

    doc.add_paragraph(f"Document type: {doc_type}")
    if situation:
        doc.add_paragraph(f"Context: {situation}")

    doc.add_heading("Summary", 1)
    red_n   = sum(1 for f in findings if f["tier"] == "red")
    amber_n = sum(1 for f in findings if f["tier"] == "amber")
    green_n = sum(1 for f in findings if f["tier"] == "green")
    doc.add_paragraph(
        f"This report identified {red_n} lawful requirement(s) not met (Red), "
        f"{amber_n} best practice gap(s) (Amber), and {green_n} compliant area(s) (Green)."
    )
    if commitments:
        doc.add_heading("Your child's EHCP commits to:", 1)

        for commitment in commitments[:20]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(commitment)
    needs_log = any(f.get("delivery_log_required") for f in findings)
    if needs_log:
        p = doc.add_paragraph(
            "⚑  One or more findings require a delivery log. "
            "If delivery of provision is not logged, it did not happen. "
            "Request the school's delivery records immediately."
        )
        p.runs[0].bold = True

    doc.add_heading("Findings", 1)

    tier_labels = {"red": "RED — Lawful Requirement Not Met",
                   "amber": "AMBER — Best Practice Gap",
                   "green": "GREEN — Compliant"}
    tier_colors = {"red": RGBColor(0xC0, 0x39, 0x2B),
                    "amber": RGBColor(0xD4, 0xA0, 0x17),
                    "green": RGBColor(0x1E, 0x84, 0x49)}

    for finding in findings:
        tier = finding["tier"]
        h = doc.add_heading(tier_labels.get(tier, tier.upper()), 2)
        h.runs[0].font.color.rgb = tier_colors.get(tier, RGBColor(0, 0, 0))

        doc.add_heading(finding["title"], 3)

        if finding.get("extract"):
            doc.add_paragraph("From your EHCP:", style="Intense Quote")
            p = doc.add_paragraph(f'"{finding["extract"][:400]}"')
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph(finding["commentary"])

        if finding.get("delivery_log_required"):
            p = doc.add_paragraph(
                "Delivery log required: If this provision has been delivered, "
                "there must be a contemporaneous record. Request the delivery log."
            )
            p.runs[0].bold = True

        doc.add_paragraph("")

    doc.add_heading("What next?", 1)
    doc.add_paragraph(
        "Red findings must be addressed at annual review. "
        "Amber findings are worth raising. "
        "Green entries are compliant — use these as benchmarks when challenging non-compliant provision. "
        "\n\nThis report is produced by FRED — Families' Rights and Entitlements Directory. "
        "It provides lawful analysis, not legal advice."
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_pdf_report(
    findings,
    child_name="your child",
    situation="",
    doc_type="EHCP",
    commitments=None
):
    """Generate PDF report using ReportLab."""
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2.5*cm, leftMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        title="FRED Report",
        author="FRED — Families' Rights and Entitlements Directory",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'fred_title', fontSize=28, alignment=TA_CENTER,
        textColor=colors.HexColor('#1a2744'),
        spaceAfter=4, spaceBefore=0,
        fontName='Helvetica-Bold',
        leading=34,
    )
    sub_style = ParagraphStyle(
        'fred_sub', fontSize=11, alignment=TA_CENTER,
        textColor=colors.HexColor('#4a5a7a'),
        spaceAfter=6, spaceBefore=0,
        fontName='Helvetica',
        leading=16,
    )
    meta_style = ParagraphStyle(
        'fred_meta', fontSize=10, alignment=TA_CENTER,
        textColor=colors.HexColor('#6a7a90'),
        spaceAfter=20, spaceBefore=0,
        fontName='Helvetica-Oblique',
        leading=14,
    )
    section_h_style = ParagraphStyle(
        'fred_section', fontSize=15, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a2744'),
        spaceAfter=8, spaceBefore=20,
        leading=20,
        borderPad=0,
    )
    h2_red = ParagraphStyle(
        'fred_h2red', fontSize=12, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#C0392B'),
        spaceAfter=4, spaceBefore=16, leading=16,
    )
    h2_amber = ParagraphStyle(
        'fred_h2amber', fontSize=12, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#D4A017'),
        spaceAfter=4, spaceBefore=16, leading=16,
    )
    h2_green = ParagraphStyle(
        'fred_h2green', fontSize=12, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1E8449'),
        spaceAfter=4, spaceBefore=16, leading=16,
    )
    finding_title_style = ParagraphStyle(
        'fred_finding_title', fontSize=11, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a1a2e'),
        spaceAfter=6, spaceBefore=4, leading=15,
    )
    extract_style = ParagraphStyle(
        'fred_extract', fontSize=10, fontName='Helvetica-Oblique',
        textColor=colors.HexColor('#555555'),
        spaceAfter=8, spaceBefore=2, leading=14,
        leftIndent=12, rightIndent=12,
    )
    body_style = ParagraphStyle(
        'fred_body', fontSize=10, fontName='Helvetica',
        textColor=colors.HexColor('#1a1a2e'),
        spaceAfter=10, spaceBefore=0, leading=15,
    )
    bold_style = ParagraphStyle(
        'fred_bold', fontSize=10, fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a1a2e'),
        spaceAfter=8, spaceBefore=4, leading=14,
    )
    footer_style = ParagraphStyle(
        'fred_footer', fontSize=9, fontName='Helvetica-Oblique',
        textColor=colors.HexColor('#888888'),
        spaceAfter=0, spaceBefore=8, leading=13, alignment=TA_CENTER,
    )

    story = []

    # ── Cover block ───────────────────────────────────────────────────────
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("FRED", title_style))
    story.append(Paragraph("Families' Rights and Entitlements Directory", sub_style))
    story.append(Spacer(1, 0.3*cm))
    # REPAIR 1: colour= → color= in all HRFlowable calls
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#1a2744'), spaceAfter=8))
    story.append(Paragraph("EHCP Analysis Report", meta_style))
    story.append(Paragraph(f"Document type: {doc_type}", meta_style))
    if situation:
        story.append(Paragraph(f"Context: {situation}", meta_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#d0dae8'), spaceAfter=16))
    story.append(Spacer(1, 0.5*cm))

    # ── Summary ───────────────────────────────────────────────────────────
    story.append(Paragraph("Summary", section_h_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#d0dae8'), spaceAfter=10))

    red_n   = sum(1 for f in findings if f["tier"] == "red")
    amber_n = sum(1 for f in findings if f["tier"] == "amber")
    green_n = sum(1 for f in findings if f["tier"] == "green")

    story.append(Paragraph(
        f'<font color="#C0392B"><b>{red_n} Red finding{"s" if red_n != 1 else ""}</b></font> — '
        f'lawful requirement{"s" if red_n != 1 else ""} not met. Must be addressed at annual review.',
        body_style
    ))
    story.append(Paragraph(
        f'<font color="#D4A017"><b>{amber_n} Amber finding{"s" if amber_n != 1 else ""}</b></font> — '
        f'best practice gap{"s" if amber_n != 1 else ""}. Worth raising at annual review.',
        body_style
    ))
    story.append(Paragraph(
        f'<font color="#1E8449"><b>{green_n} Green finding{"s" if green_n != 1 else ""}</b></font> — '
        f'compliant. Use as benchmark when challenging non-compliant provision.',
        body_style
    ))

    story.append(Spacer(1, 0.4*cm))

    needs_log = any(f.get("delivery_log_required") for f in findings)
    if needs_log:
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#D4A017'), spaceAfter=6))
        story.append(Paragraph(
            "⚑  Delivery log required",
            bold_style
        ))
        story.append(Paragraph(
            "One or more findings require a delivery log. "
            "If provision has been delivered, there must be a contemporaneous record. "
            "If it is not logged, it did not happen. "
            "Request the school's delivery records immediately.",
            body_style
        ))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#D4A017'), spaceAfter=10))
        
        if commitments:
            story.append(Paragraph("Your child's EHCP commits to:", section_h_style))

        for commitment in commitments[:20]:
            safe_commitment = (
                commitment.replace("&", "&amp;")
                          .replace("<", "&lt;")
                          .replace(">", "&gt;")
            )
            story.append(Paragraph(f"• {safe_commitment}", body_style))

        story.append(Spacer(1, 0.3*cm))
   
    story.append(Spacer(1, 0.6*cm))

    # ── Findings ──────────────────────────────────────────────────────────
    story.append(Paragraph("Findings", section_h_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#d0dae8'), spaceAfter=12))

    tier_h_styles = {"red": h2_red, "amber": h2_amber, "green": h2_green}
    tier_labels   = {
        "red":   "RED — Lawful Requirement Not Met",
        "amber": "AMBER — Best Practice Gap",
        "green": "GREEN — Compliant",
    }

    for i, finding in enumerate(findings):
        tier = finding["tier"]
        story.append(Paragraph(tier_labels.get(tier, tier.upper()), tier_h_styles.get(tier, body_style)))
        story.append(Paragraph(finding['title'], finding_title_style))

        if finding.get("extract"):
            safe_extract = finding["extract"][:400].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(
                "From your EHCP:",
                ParagraphStyle(
                    'from_ehcp_label',
                    fontSize=8,
                    fontName='Helvetica-Bold',
                    textColor=colors.HexColor('#888888'),
                    spaceAfter=2,
                    spaceBefore=6,
                    leading=11,
                    leftIndent=12,
                )
            ))
            story.append(Paragraph(f'"{safe_extract}"', extract_style))

        commentary_parts = finding["commentary"].split("\n\n")
        for part in commentary_parts:
            if part.strip():
                story.append(Paragraph(part.strip(), body_style))

        if finding.get("delivery_log_required"):
            story.append(Paragraph(
                "⚑  Delivery log required: if this provision has been delivered, "
                "there must be a contemporaneous record. Request the delivery log.",
                bold_style
            ))

        story.append(Spacer(1, 0.3*cm))
        story.append(HRFlowable(width="100%", thickness=0.3, color=colors.HexColor('#e0e8f0'), spaceAfter=8))

    story.append(Spacer(1, 0.8*cm))

    # ── What next ─────────────────────────────────────────────────────────
    story.append(Paragraph("What next?", section_h_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#d0dae8'), spaceAfter=12))

    story.append(Paragraph(
        "<b>Red findings</b> must be addressed at your annual review. "
        "Where there are shortcomings in the delivery of specified provision, "
        "the school has a duty of care — the duty to deliver what is written in the EHCP "
        "is absolute under the Children and Families Act 2014. It is not discretionary and "
        "it does not depend on resources. If provision has not been delivered, you have the "
        "right to ask for evidence of delivery and to request that the annual review formally "
        "records the gap.",
        body_style
    ))
    story.append(Paragraph(
        "<b>Amber findings</b> are worth raising. They are not lawful failures but they affect "
        "the quality and accountability of your child's support.",
        body_style
    ))
    story.append(Paragraph(
        "<b>Green entries</b> are your benchmarks. Use them when challenging non-compliant provision — "
        "if one section of the EHCP meets the standard, there is no reason another cannot.",
        body_style
    ))

    story.append(Spacer(1, 0.8*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1a2744'), spaceAfter=8))
    story.append(Paragraph(
        "This report is produced by FRED — Families' Rights and Entitlements Directory. "
        "It provides lawful analysis, not legal advice. "
        "Reference: Children and Families Act 2014 and SEND Code of Practice 2015.",
        footer_style
    ))

    doc.build(story)
    buf.seek(0)
    return buf


# ── SESSION STATE DEFAULTS ────────────────────────────────────────────────────

def init_state():
    defaults = {
        "stage": "landing",
        "findings": [],
        "provision_inventory": {},
        "parse_meta": {},
        "full_text": "",
        "doc_type": "Final EHCP",
        "situation": "",
        "email_submitted": False,
        "email_address": "",
        "survey_submitted": False,
        "show_full_report": False,
        "relationship_tone": "neutral",
        "upcoming_dates": "",
        "doc_name": "",
        "subscribed": False,
        "corr_analysed": False,
        "show_companion": False,
        "show_amendment": False,
        "show_all_patterns": False,
        "tone_override": None,
        # REPAIR 3: input_method persists across reruns
        "input_method": "Upload file (PDF or Word)",
        # Stored paste text — survives rerun
        "paste_text_stored": "",
        "vault": {},
        "thread": [],
        "thread_context": "",
        "ehc_request_started": False,
        "ehc_journey_active": False,
        "ehc_current_category": 1,"ehc_request_started": False,
        "ehc_journey_active": False,
        "ehc_current_category": 1,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


# ── HELPER RENDERERS ──────────────────────────────────────────────────────────

def _highlight_trigger(text: str, trigger: str) -> str:
    """Bolds the trigger word in red for HTML display. Used in render_finding_card only."""
    if not trigger or not text:
        return text
    return re.sub(
        rf'({re.escape(trigger)})',
        r'<b style="color:#C0392B">\1</b>',
        text,
        flags=re.IGNORECASE
    )
def render_finding_card(finding, index=None, show_full=True):
    tier = finding["tier"]
    badge_class = f"badge-{tier}"
    card_class  = f"finding-{tier}"
    tier_labels = {
        "red":   "RED — Lawful Requirement Not Met",
        "amber": "AMBER — Best Practice Gap",
        "green": "GREEN — Compliant",
    }
    label = tier_labels.get(tier, tier.upper())

    html = f"""
    <div class="{card_class}">
      <span class="{badge_class}">{label}</span>
      <p style="font-weight:700;margin:0.4rem 0 0.5rem 0;font-size:1rem;">{finding['title']}</p>
    """

    extract = finding.get("extract") or ""
    if extract and show_full:
        html += (
            '<div style="background:#f5f5f0;border-left:3px solid #ccc;'
            'border-radius:0 4px 4px 0;padding:0.6rem 1rem;margin:0.5rem 0 0.8rem 0;">'
            '<p style="font-size:0.78rem;font-weight:600;text-transform:uppercase;'
            'letter-spacing:0.08em;color:#888;margin:0 0 0.3rem 0;">From your EHCP:</p>'
            '<p style="font-style:italic;color:#333;font-size:0.92rem;'
            'margin:0;line-height:1.6;">'
            f'"{_highlight_trigger(extract[:800], finding.get("trigger_word", ""))}"'
            '</p></div>'
        )

    if show_full:
        html += f'<p style="margin:0;font-size:0.95rem;line-height:1.6;">{finding["commentary"]}</p>'

    if finding.get("delivery_log_required") and show_full:
        html += '<p style="font-weight:700;margin-top:0.7rem;font-size:0.9rem;">⚑ Delivery log required</p>'

    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)


def render_traffic_light_explainer():
    st.markdown("### How FRED grades your report")
    cols = st.columns(3)
    with cols[0]:
        st.markdown(f"""
        <div style="border-left:5px solid {RED};padding:1rem;background:#fdf4f3;border-radius:0 6px 6px 0;">
          <span class="badge-red">RED</span>
          <p style="font-weight:700;margin:0.4rem 0 0.2rem;">Lawful requirement not met</p>
          <p style="font-size:0.9rem;margin:0;">Must be addressed at annual review. Not lawfully enforceable under the Children and Families Act 2014 as written.</p>
        </div>
        """, unsafe_allow_html=True)
    with cols[1]:
        st.markdown(f"""
        <div style="border-left:5px solid {AMBER};padding:1rem;background:#fdf9f0;border-radius:0 6px 6px 0;">
          <span class="badge-amber">AMBER</span>
          <p style="font-weight:700;margin:0.4rem 0 0.2rem;">Best practice gap</p>
          <p style="font-size:0.9rem;margin:0;">Worth raising at annual review. Not a lawful failure but a gap in best practice that may affect your child.</p>
        </div>
        """, unsafe_allow_html=True)
    with cols[2]:
        st.markdown(f"""
        <div style="border-left:5px solid {GREEN};padding:1rem;background:#f3faf5;border-radius:0 6px 6px 0;">
          <span class="badge-green">GREEN</span>
          <p style="font-weight:700;margin:0.4rem 0 0.2rem;">Compliant</p>
          <p style="font-size:0.9rem;margin:0;">Provision meets lawful requirements. Use compliant entries as benchmarks when challenging non-compliant areas.</p>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)


# ── PAGES ─────────────────────────────────────────────────────────────────────

def page_landing():
    st.markdown("""
    <div class="hero">
      <h1>FRED</h1>
      <p class="full-name">Families' Rights and Entitlements Directory</p>
      <p class="tagline">Know what your child is entitled to.<br>Know when it's not being delivered.</p>
      <p class="origin">Built by a parent who learned the hard way — so you don't have to.</p>
      <p class="sub">
        FRED analyses your child's EHCP against the Children and Families Act 2014 and the SEND Code of Practice.
        Plain-English findings. Red, amber, green. No jargon.
      </p>
      <p class="service-line">One-off report · Annual subscription · Monthly plan</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns([1, 2, 2, 1])
    with col2:
        if st.button("Analyse my documents", use_container_width=True, key="landing_analyse"):
            st.session_state.stage = "explainer"
            st.rerun()
    with col3:
        if st.button("Apply for an EHCP", use_container_width=True, key="landing_ehc"):
            st.session_state.stage = "ehc_request"
            st.rerun()

    st.markdown("""
    <p style="text-align: center;color: #6a7a90;font-size:0.9rem;margin-top:0.8rem;">
      Upload first. Decide after. Your report is ready before you pay.
    </p>
    <p style="text-align: center;color: #8a9ab0;font-size:0.82rem;">
      From £XX for the full report — or see our subscription plans below
    </p>
    """, unsafe_allow_html=True)

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    st.markdown("<h2 style='text-align: center;'>Everything you need to know.</h2>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    hiw_items = [
        "Gather whatever documents you have — your child's EHCP, EP report, OT report, school correspondence. You don't need everything.",
        "Upload your EHCP. FRED reads it against the Children and Families Act 2014 and the SEND Code of Practice.",
        "FRED gives you one real finding immediately, at no cost. You'll see exactly what kind of analysis you're getting before you decide anything.",
        "Your full report is ready. Red findings are lawful requirements not met. Amber are best practice gaps. Green are compliant — use these as your benchmarks.",
        "Download your report as a Word or PDF document. Take it to your annual review. Use it in correspondence. It's yours.",
        "A subscription gives you the full intelligence layer — correspondence analysis, email drafting, meeting preparation, document vault, and more.",
    ]

    for item in hiw_items:
        st.markdown(f"""
        <div class="hiw-item">
          <div class="hiw-dot"></div>
          <div style="font-size:0.97rem;line-height:1.6;">{item}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    render_traffic_light_explainer()

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    st.markdown("<h2 style='text-align: center;'>Plans</h2>", unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    p1, p2, p3 = st.columns(3)
    with p1:
        st.markdown(f"""
        <div class="pricing-card">
          <p style="font-weight:700;font-size:1.1rem;margin-bottom:0.3rem;">One-off Report</p>
          <p class="price">£XX</p>
          <p class="period">single purchase, no commitment</p>
          <ul>
            <li>Full EHCP analysis</li>
            <li>Red / Amber / Green findings</li>
            <li>Word and PDF download</li>
            <li>No ongoing subscription</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)

    with p2:
        st.markdown(f"""
        <div class="pricing-card featured">
          <span class="best-value-tag">Best value</span><br>
          <p style="font-weight:700;font-size:1.1rem;margin-bottom:0.3rem;">Annual Subscription</p>
          <p class="price">£XX</p>
          <p class="period">per year — includes first report</p>
          <ul>
            <li>Full EHCP analysis</li>
            <li>Document vault</li>
            <li>Correspondence analysis</li>
            <li>Email drafting</li>
            <li>Meeting preparation</li>
            <li>Annual review preparation</li>
            <li>Reduced renewal rate year two</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)

    with p3:
        st.markdown(f"""
        <div class="pricing-card">
          <p style="font-weight:700;font-size:1.1rem;margin-bottom:0.3rem;">Monthly</p>
          <p class="price">£XX</p>
          <p class="period">first month includes report, then £XX/month</p>
          <ul>
            <li>Full EHCP analysis</li>
            <li>Document vault</li>
            <li>Correspondence analysis</li>
            <li>Email drafting</li>
            <li>Cancel anytime</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><p style='text-align: center;color: #6a7a90;font-size:0.85rem;'>No hidden charges. Your report is ready before you purchase.</p>", unsafe_allow_html=True)

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    st.markdown("<h2 style='text-align: center;'>FAQ</h2>", unsafe_allow_html=True)

    faqs = [
        ("When do I pay?", "After you see one real finding from your EHCP, before you access the full report. You upload first — your report is ready before you pay."),
        ("Is this legal advice?", "No. FRED provides lawful analysis — it tells you what the law requires. It does not tell you what to do about it. That decision is yours."),
        ("What documents do I need?", "Upload whatever you have. An EP report, OT report, school correspondence, diagnosis letter — FRED will work with any of these. If you don't have an EHCP yet, upload what you do have and FRED will advise on next steps to ensure the right support is in place for your child."),
        ("What is an EHCP?", "An Education, Health and Care Plan is a legally binding document setting out a child's needs and the provision that must be made. The word 'must' in that sentence is important."),
        ("What if I only have a draft EHCP?", "FRED analyses draft and final EHCPs differently. For a draft, FRED highlights what should be strengthened before the document is finalised. For a final EHCP, FRED references your rights and the duty to deliver."),
        ("Is my data private?", "Yes. Your documents are used only for your analysis. They are not shared, stored beyond your session, or used to train any system."),
    ]

    for q, a in faqs:
        with st.expander(q):
            st.write(a)

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)

    st.markdown("### Want to shape what FRED becomes?")
    st.markdown("We're in beta. Leave your thoughts — takes two minutes.")
    st.markdown(f"""
    <a href="{GOOGLE_FORM_URL}" target="_blank" style="
        display:inline-block;
        background:#1a2744;
        color: white;
        padding:0.7rem 1.8rem;
        border-radius:4px;
        text-decoration:none;
        font-family:'Source Sans 3',sans-serif;
        font-weight:600;
        font-size:0.95rem;
    ">Share your thoughts →</a>
    <p style="font-size:0.8rem;color: #888;margin-top:0.5rem;">Opens in a new tab. No obligation.</p>
    """, unsafe_allow_html=True)


def page_explainer():
    st.markdown("## Before you upload")
    render_traffic_light_explainer()

    st.markdown("---")
    st.markdown("### What would you like to analyse?")

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("Continue to upload", use_container_width=True):
            st.session_state.stage = "upload"
            st.rerun()


def page_upload():
    st.markdown("## Upload your documents")

    st.markdown(f"""
    <div style="background:#eaf5e0;border-radius:10px;padding:1rem 1.3rem;margin-bottom:1.5rem;border:0.5px solid #c0ddb0;">
      <p style="margin:0;font-size:0.95rem;color: #2d4a2d;line-height:1.7;">
        Upload your child's EHCP first. Then add any other documents you have —
        FRED will read them all and refer back to them when you analyse correspondence
        or request an amendment.
      </p>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Analyse my documents", use_container_width=False, key="analyse_top2"):
        st.session_state["analyse_clicked_top"] = True
    DOCUMENT_TYPES = [
        {
            "key": "ehcp",
            "label": "EHCP",
            "description": "Essential. FRED analyses Section F provision, Section E outcomes, and Section B needs. Everything else is cross-referenced against this.",
            "required": False,
        },
        {
            "key": "ep_report",
            "label": "EP report",
            "description": "FRED checks whether the EP's recommendations have been converted into specified provision in Section F — or whether they have been laundered into vague language.",
            "required": False,
        },
        {
            "key": "ot_report",
            "label": "OT report",
            "description": "FRED checks whether OT recommendations appear in Section F and whether the OT provision loop has been closed — treatment block, discharge, re-referral threshold.",
            "required": False,
        },
        {
            "key": "sen_policy",
            "label": "School SEN policy",
            "description": "FRED cross-references the school's own SEN policy against what is in the EHCP and what the correspondence shows. The school cannot dispute its own policy.",
            "required": False,
        },
        {
            "key": "accessibility_plan",
            "label": "School accessibility plan",
            "description": "FRED checks for unfulfilled commitments — acoustic surveys, staff training, environmental audits — and flags gaps between policy commitment and practice.",
            "required": False,
        },
        {
            "key": "diagnosis",
            "label": "Diagnosis letter",
            "description": "FRED uses this to confirm the need basis for provision and to check whether the diagnosis profile is accurately reflected in Section B.",
            "required": False,
        },
    ]

    vault = st.session_state.vault

    for doc in DOCUMENT_TYPES:
        key  = doc["key"]
        already_uploaded = key in vault

        col1, col2 = st.columns([2, 3])
        with col1:
            label = f"{doc['label']} {'✓' if already_uploaded else '— required' if doc['required'] else '— optional'}"
            uploaded = st.file_uploader(
                label,
                type=["pdf", "docx", "txt"],
                key=f"vault_{key}",
                help=doc["description"],
            )
            if uploaded:
                text = extract_text_from_pdf(uploaded) if uploaded.name.endswith(".pdf") else extract_text_from_docx(uploaded)
                vault[key] = {"name": uploaded.name, "text": text}
                st.session_state.vault = vault

        with col2:
            if already_uploaded or uploaded:
                st.markdown(f"""
                <div style="background:#eaf5e0;border-radius:8px;padding:0.6rem 1rem;margin-top:1.6rem;border:0.5px solid #c0ddb0;">
                  <p style="margin:0;font-size:0.85rem;color: #2d5a2d;font-weight:500;">
                    ✓ FRED has read the {doc['label']}
                  </p>
                  <p style="margin:0.2rem 0 0;font-size:0.8rem;color: #5a8a5a;">{doc['description']}</p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="padding:0.6rem 1rem;margin-top:1.6rem;">
                  <p style="margin:0;font-size:0.82rem;color: #888;">{doc['description']}</p>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("---")

    if vault:
        loaded = ", ".join(d["label"] for d in DOCUMENT_TYPES if d["key"] in vault)
        st.markdown(f"""
        <div style="background:#2d4a2d;border-radius:10px;padding:0.9rem 1.2rem;margin-bottom:1rem;">
          <p style="margin:0;font-size:0.88rem;color: #9dc98a;font-weight:500;">
            FRED is holding: {loaded}
          </p>
          <p style="margin:0.3rem 0 0;font-size:0.82rem;color: #7ab870;">
            These will be cross-referenced when you analyse correspondence or request an amendment.
          </p>
        </div>
        """, unsafe_allow_html=True)

    col_a, col_b = st.columns([1, 4])
    with col_a:
        if st.button("Clear vault — start again", key="clear_vault"):
            st.session_state.vault = {}
            st.rerun()
    with col_b:
        if vault:
            st.markdown(
                "<p style='font-size:0.82rem;color: #888;padding-top:0.4rem;'>"
                "To add more documents, upload them above. To replace a document, "
                "upload a new version — it will overwrite the previous one."
                "</p>",
                unsafe_allow_html=True
            )

    analyse_clicked = st.button("Analyse my documents", use_container_width=False, key="analyse_top") or st.session_state.get("analyse_clicked_top", False)
    st.session_state["analyse_clicked_top"] = False

    st.markdown("---")
    st.markdown("### A few quick questions")
    st.markdown("<p style='font-size:0.88rem;color: #666;margin-bottom:1rem;'>These help FRED tailor the analysis. Nothing here is mandatory.</p>", unsafe_allow_html=True)

    doc_type = st.radio(
        "What is the primary document you have uploaded?",
        ["Draft EHCP", "Final EHCP", "EP report", "OT report", "School correspondence", "Other"],
        index=1,
        horizontal=True,
    )

    situation = st.text_area(
        "Briefly describe your current situation (optional)",
        placeholder="e.g. Annual review coming up in March, school not delivering speech therapy, considering appeal...",
        height=80,
    )

    upcoming_dates = st.text_input(
        "Any upcoming dates we should know about? (optional)",
        placeholder="e.g. Annual review 15 March, appeal deadline 1 April"
    )

    relationship_tone = st.select_slider(
        "How would you describe your current relationship with the school?",
        options=["Very difficult", "Difficult", "Neutral", "Generally positive", "Very positive"],
        value="Neutral",
        help="We use this to help shape the tone of any suggested emails."
    )

    st.markdown("---")

    if analyse_clicked:
        if not vault:
            st.error("Please upload at least one document to continue.")
        else:
            with st.spinner("Reading your documents…"):
                combined_text = " ".join(v["text"] for v in vault.values())
                findings, meta = run_full_analysis(combined_text)
                if INVENTORY_AVAILABLE:
                    inventory = build_provision_inventory(combined_text)
                    st.session_state.provision_inventory = inventory

            st.session_state.findings          = findings
            st.session_state.parse_meta        = meta
            st.session_state.full_text         = combined_text
            st.session_state.doc_type          = doc_type
            st.session_state.situation         = situation
            st.session_state.upcoming_dates    = upcoming_dates
            st.session_state.relationship_tone = relationship_tone
            st.session_state.doc_name          = vault.get("ehcp", {}).get("name", "EHCP")
            write_success = write_analysis_to_supabase(findings, meta)
            st.session_state["analysis_saved"] = write_success
            st.session_state.stage             = "sneak_peek"
            st.rerun()



def page_sneak_peek():
    # Logged in users bypass the email gate — beta access
    if st.session_state.get("user"):
        st.session_state.email_submitted = True
        st.session_state.email_address = st.session_state["user"].email
        st.session_state.stage = "full_report"
        st.rerun()

    findings = st.session_state.findings

    st.markdown("## Your report is ready")
    st.markdown(
        "Here is one real finding from your EHCP — exactly what the full report looks like. "
        "No blur. No teaser. The most significant finding from your document."
    )

    red_findings = [f for f in findings if f["tier"] == "red"]
    preview_finding = red_findings[0] if red_findings else findings[0] if findings else None

    if preview_finding:
        st.markdown('<div class="sneak-box">', unsafe_allow_html=True)
        st.markdown("**Preview — from your EHCP:**")
        render_finding_card(preview_finding, show_full=True)
        st.markdown('</div>', unsafe_allow_html=True)

    red_n   = sum(1 for f in findings if f["tier"] == "red")
    amber_n = sum(1 for f in findings if f["tier"] == "amber")
    green_n = sum(1 for f in findings if f["tier"] == "green")

    st.markdown(f"""
    <div style="background:white;border:1px solid #d0dae8;border-radius:8px;padding:1.2rem;margin:1.2rem 0;">
      <p style="margin:0;font-size:1rem;">
        Your full report contains:&nbsp;
        <span style="color:{RED};font-weight:700;">{red_n} Red</span> &nbsp;·&nbsp;
        <span style="color:{AMBER};font-weight:700;">{amber_n} Amber</span> &nbsp;·&nbsp;
        <span style="color:{GREEN};font-weight:700;">{green_n} Green</span>
      </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    # ── Email gate ────────────────────────────────────────────────────────
    if not st.session_state.email_submitted:
        st.markdown("### Access your full report")
        st.markdown(
            "Enter your email address to access your full report. "
            "During beta, everything is free — no card details, no obligation."
        )
        st.markdown(f"""
        <div style="background:{NAVY};border-radius:8px;padding:1.5rem 2rem;margin:1rem 0;">
          <p style="color:#a8b8d8;font-size:0.85rem;margin:0 0 0.3rem;">
            FRED BETA ACCESS
          </p>
          <p style="color:white;font-size:1.1rem;font-weight:600;margin:0 0 1rem;">
            Your report is waiting. Enter your email to continue.
          </p>
        </div>
        """, unsafe_allow_html=True)

        with st.form("email_gate"):
            email = st.text_input(
                "Email address",
                placeholder="your@email.com",
            )
            submitted = st.form_submit_button("Access my report →", use_container_width=True)
            if submitted:
                if email and "@" in email and "." in email:
                    st.session_state.email_submitted = True
                    st.session_state.email_address   = email
                    print(f"FRED BETA SIGNUP: {email} — doc: {st.session_state.doc_name}")
                    st.rerun()
                else:
                    st.error("Please enter a valid email address.")
    else:
        st.session_state.stage = "full_report"
        st.rerun()


def page_full_report():
    findings  = st.session_state.findings
    doc_type  = st.session_state.doc_type
    situation = st.session_state.situation
    meta      = st.session_state.parse_meta

    red_n   = sum(1 for f in findings if f["tier"] == "red")
    amber_n = sum(1 for f in findings if f["tier"] == "amber")
    green_n = sum(1 for f in findings if f["tier"] == "green")

    st.markdown("## Your full FRED report")
    if st.session_state.get("analysis_saved"):
        st.caption("Your analysis has been saved to your account.")
        
    st.markdown(f"""
    <div style="background:#f0f4fa;border-radius:6px;padding:0.8rem 1.2rem;margin-bottom:1.2rem;font-size:0.9rem;color: #444;">
      Document type: <b>{doc_type}</b> &nbsp;·&nbsp;
      Section F blocks found: <b>{meta.get('f_blocks_found', 0)}</b> &nbsp;·&nbsp;
      Section E blocks found: <b>{meta.get('e_blocks_found', 0)}</b>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div style="background:white;border:1px solid #d0dae8;border-radius:8px;padding:1.2rem;margin-bottom:1.5rem;">
      <p style="margin:0;font-size:1.05rem;font-weight:600;">Summary</p>
      <p style="margin:0.5rem 0 0;font-size:0.97rem;">
        <span style="color:{RED};font-weight:700;">{red_n} Red</span> — lawful requirement(s) not met.
        Must be addressed at annual review.<br>
        <span style="color:{AMBER};font-weight:700;">{amber_n} Amber</span> — best practice gap(s).
        Worth raising at annual review.<br>
        <span style="color:{GREEN};font-weight:700;">{green_n} Green</span> — compliant.
        Use as benchmark when challenging non-compliant provision.
      </p>
    </div>
    """, unsafe_allow_html=True)

    needs_log = any(f.get("delivery_log_required") for f in findings)
    if needs_log:
        st.markdown(f"""
        <div style="background:#fff8e1;border:1px solid {AMBER};border-radius:6px;padding:0.9rem 1.2rem;margin-bottom:1.2rem;">
          <p style="margin:0;font-weight:700;">⚑  Delivery log required</p>
          <p style="margin:0.4rem 0 0;font-size:0.9rem;">
            One or more findings require a delivery log.
            If provision has been delivered, there must be a contemporaneous record.
            If not logged, it did not happen.
            Request the school's delivery records immediately.
          </p>
        </div>
        """, unsafe_allow_html=True)

    if "Draft" in doc_type:
        st.info(
            "This is a draft EHCP. Red findings indicate language and gaps that should be "
            "strengthened before the document is finalised. You have a window to request amendments."
        )

# ── EHCP Commitments Summary ──────────────────────────────────────────────
    f_blocks_for_summary = find_section_blocks(st.session_state.get("full_text", ""), "F")
    commitments = build_ehcp_commitments_summary(f_blocks_for_summary)

    if commitments:
        st.markdown("### Your child's EHCP commits to:")
        st.markdown(
            "<p style='font-size:0.88rem;color:#666;margin-bottom:0.8rem;'>"
            "These are the specific commitments written into Section F of the plan. "
            "Each one is a legal obligation on the school and the LA."
            "</p>",
            unsafe_allow_html=True
        )
        for item in commitments:
            st.markdown(f"- {item}")
        st.markdown("---")
    st.markdown("---")
    if "provision_inventory" in st.session_state:

     inventory = st.session_state.provision_inventory

    if inventory:

        st.markdown("## Provision Inventory")

        st.markdown(
            """
            These are the approaches your child's documents identify as useful,
            recommended, or currently referenced.

            This is not a complete delivery record.

            At your next review, ask the school which of these are currently in place,
            how they are being delivered, and how impact is being monitored.
            """
        )

        for category, items in inventory.items():

            st.markdown(f"### {category}")

            for item in items:
                st.markdown(f"- {item}")
    st.markdown("### Findings")

    for i, finding in enumerate(findings):
        if finding["tier"] == "green" and green_n == 0:
            continue
        render_finding_card(finding, index=i, show_full=True)

    st.markdown("---")
    st.subheader("Professional Reports — Recommendation Check")
    st.caption(
        "Upload an EP, SALT, or OT report. FRED will check whether each "
        "recommendation has been written into Section F as a legal commitment."
    )

    report_type = st.selectbox(
        "Report type",
        ["EP report", "SALT report", "OT report", "Other professional report"],
        key="report_type_select"
    )

    uploaded_report = st.file_uploader(
        f"Upload {report_type} (PDF)",
        type=["pdf"],
        key="report_uploader"
    )

    if uploaded_report is not None:
        with st.spinner(f"Reading {report_type}..."):
            report_text = extract_text_from_pdf(uploaded_report)

        if not report_text or len(report_text.strip()) < 50:
            st.error(
                f"Could not extract text from this {report_type}. "
                "Check the PDF is not scanned-only."
            )
        else:
            section_f_blocks = find_section_blocks(
                st.session_state.get("full_text", ""), "F"
            )
            section_f_combined = "\n".join(section_f_blocks) if section_f_blocks else ""

            if not section_f_combined:
                st.warning(
                    "Section F has not been extracted yet. "
                    "Upload your EHCP first, then upload the report."
                )
            else:
                rec_findings = analyse_report_recommendations(
                    report_text=report_text,
                    section_f_text=section_f_combined,
                    report_label=report_type
                )

                if not rec_findings:
                    st.info(
                        f"No recommendation language detected in this {report_type}. "
                        "Check the report extracted correctly using the text preview below."
                    )
                    with st.expander("Report text preview (first 1000 characters)"):
                        st.text(report_text[:1000])
                else:
                    red_count   = sum(1 for f in rec_findings if f["type"] == "RED")
                    green_count = sum(1 for f in rec_findings if f["type"] == "GREEN")

                    st.markdown(
                        f"**{report_type} — Recommendation check:** "
                        f":red[{red_count} not converted] &nbsp;|&nbsp; "
                        f":green[{green_count} converted to Section F]"
                    )

                    for finding in rec_findings:
                        if finding["type"] == "RED":
                            with st.container(border=True):
                                st.markdown("🔴 **Recommendation not converted to Section F**")
                                st.markdown("**From the report:**")
                                st.markdown(f"*{finding['recommendation'][:400]}*")
                                st.markdown(f"**Trigger phrase found:** `{finding['trigger']}`")
                                st.markdown(
                                    f"**Substance checked for:** "
                                    f"{', '.join(finding['key_terms'])}"
                                )
                                st.markdown(f"**Finding:** {finding['finding']}")
                                st.caption(f"📋 {finding['citation']}")

                        elif finding["type"] == "GREEN":
                            with st.expander(
                                f"✅ Converted — {', '.join(finding['matched_terms'][:3])}"
                            ):
                                st.markdown("**From the report:**")
                                st.markdown(f"*{finding['recommendation'][:400]}*")
                                st.markdown("**Matched in Section F:**")
                                st.markdown(f"*{finding['section_f_match'][:400]}*")
                                st.markdown(f"{finding['finding']}")
                                
    st.markdown("### What next?")
    st.markdown(
        "**Red findings** must be addressed at your annual review. "
        "Where there are shortcomings in the delivery of specified provision, "
        "the school has a duty of care — the duty to deliver what is written in the EHCP "
        "is absolute under the Children and Families Act 2014. It is not discretionary and "
        "it does not depend on resources. If provision has not been delivered, you have the "
        "right to ask for evidence of delivery and to request that the annual review formally "
        "records the gap."
        "\n\n"
        "**Amber findings** are worth raising. They are not lawful failures but they affect "
        "the quality and accountability of your child's support."
        "\n\n"
        "**Green entries** are your benchmarks. Use them when challenging non-compliant provision — "
        "if one section of the EHCP meets the standard, there is no reason another cannot."
    )

    st.markdown("---")
    if not st.session_state.subscribed:
        st.markdown(f"""
        <div style="background:{NAVY};border-radius:8px;padding:2rem;margin:1rem 0;text-align: center;">
          <p style="color:#a8b8d8;font-size:0.85rem;margin:0 0 0.4rem;letter-spacing:0.1em;text-transform:uppercase;">
            FRED BETA — FREE ACCESS
          </p>
          <p style="color:white;font-size:1.3rem;font-weight:700;margin:0 0 0.6rem;">
            Want the full FRED experience?
          </p>
          <p style="color:#c8d8f0;font-size:1rem;margin:0 0 1.2rem;max-width:480px;margin-left:auto;margin-right:auto;">
            Subscribe free during beta and unlock correspondence analysis,
            email drafting, meeting preparation, document vault, and more.
          </p>
          <p style="color:#a8b8d8;font-size:0.85rem;margin:0;">
            Already registered as {st.session_state.email_address if st.session_state.email_address else "a beta tester"}
          </p>
        </div>
        """, unsafe_allow_html=True)
        col1, col2, col3 = st.columns([1,1,1])
        with col2:
            if st.button("Subscribe free — unlock everything", use_container_width=True):
                st.session_state.subscribed = True
                st.session_state.stage = "subscriber"
                st.rerun()
    else:
        if st.button("Go to my FRED workspace →"):
            st.session_state.stage = "subscriber"
            st.rerun()

    st.markdown("---")
    st.markdown("### Help us build FRED properly")
    st.markdown(
        "Two minutes. Completely optional. "
        "Your feedback shapes what FRED becomes."
    )
    st.markdown(f"""
    <a href="{GOOGLE_FORM_URL}" target="_blank" style="
        display:inline-block;
        background:{NAVY};
        color: white;
        padding:0.7rem 1.8rem;
        border-radius:4px;
        text-decoration:none;
        font-family:'Source Sans 3',sans-serif;
        font-weight:600;
        font-size:0.95rem;
    ">Leave feedback →</a>
    <p style="font-size:0.8rem;color: #888;margin-top:0.5rem;">
        Opens in a new tab. Takes 2 minutes.
    </p>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### Download your report")

    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        word_buf = generate_word_report(
            findings,
            doc_type=doc_type,
            situation=situation,
            commitments=commitments
        )
        st.download_button(
            label="Download as Word",
            data=word_buf,
            file_name="FRED_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with col2:
        pdf_buf = generate_pdf_report(
            findings,
            doc_type=doc_type,
            situation=situation,
            commitments=commitments
        )
        st.download_button(
            label="Download as PDF",
            data=pdf_buf,
            file_name="FRED_report.pdf",
            mime="application/pdf",
        )

    st.markdown(
        "<p style='font-size:0.85rem;color: #888;'>Word is best for Windows/Office users. PDF is best for Apple devices.</p>",
        unsafe_allow_html=True
    )

    st.markdown("---")
    st.markdown("### Want the full intelligence layer?")
    st.markdown(
        "A FRED subscription gives you correspondence analysis, email drafting, "
        "meeting preparation, post-meeting summaries, document vault, and annual review preparation. "
        "Subscriptions open at launch."
    )

def page_login():
    st.markdown("""
    <style>
    [data-testid="stAppViewContainer"] {
        background-color: #f5f3ef;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="max-width:480px;margin:3rem auto 0 auto;padding:0 1rem;">
      <div style="text-align:center;margin-bottom:2rem;">
        <h1 style="font-family:'DM Serif Display',serif;font-size:3.2rem;
                   color:#2d4a2d;margin:0;letter-spacing:0.06em;">FRED</h1>
        <p style="font-size:0.72rem;letter-spacing:0.22em;text-transform:uppercase;
                  color:#5a8a5a;margin:0.3rem 0 1.2rem;font-weight:300;">
          Families' Rights and Entitlements Directory
        </p>
        <p style="font-size:1rem;color:#444;line-height:1.7;margin:0;font-weight:300;">
          FRED helps parents understand what their child's EHCP commits to
          and whether it is being delivered.
        </p>
      </div>
    </div>
    """, unsafe_allow_html=True)

    mode = st.radio(
        "What would you like to do?",
        ["Sign in", "Create account"],
        horizontal=True,
        key="login_mode"
    )

    email = st.text_input("Email address", key="login_email")
    password = st.text_input("Password", type="password", key="login_password")

    if mode == "Sign in":
        if st.button("Sign in", key="signin_btn", use_container_width=True):
            if not email or not password:
                st.error("Please enter your email and password.")
            elif not SUPABASE_AVAILABLE:
                st.error("Authentication unavailable. Please try again later.")
            else:
                try:
                    result = supabase.auth.sign_in_with_password({
                        "email": email,
                        "password": password
                    })
                    st.session_state["user"] = result.user
                    st.session_state["session"] = result.session
                    st.rerun()
                except Exception as e:
                    st.error(f"Sign in failed: {e}")

    else:
        st.markdown(
            "<p style='font-size:0.88rem;color:#666;'>Free during beta — no card required.</p>",
            unsafe_allow_html=True
        )
        if st.button("Create account", key="signup_btn", use_container_width=True):
            if not email or not password:
                st.error("Please enter your email and password.")
            elif len(password) < 6:
                st.error("Password must be at least 6 characters.")
            elif not SUPABASE_AVAILABLE:
                st.error("Authentication unavailable. Please try again later.")
            else:
                try:
                    result = supabase.auth.sign_up({
                        "email": email,
                        "password": password
                    })
                    st.session_state["user"] = result.user
                    st.session_state["session"] = result.session
                    st.success(
                        "Account created. Check your email to confirm your address, "
                        "then sign in."
                    )
                except Exception as e:
                    st.error(f"Account creation failed: {e}")

    st.markdown("""
    <p style="text-align:center;font-size:0.8rem;color:#999;margin-top:1.2rem;font-weight:300;">
      Free during beta · No card required · Built by a parent
    </p>
    """, unsafe_allow_html=True)

def page_survey():
    st.markdown("### Your feedback")
    st.markdown("This takes two minutes and helps us build FRED properly.")

    with st.form("survey_form"):
        q1 = st.radio("Did the report identify anything you didn't already know?",
                      ["Yes, something new", "Confirmed what I suspected", "Nothing new", "Not sure"])
        q2 = st.radio("Were the traffic light grades (Red / Amber / Green) clear?",
                      ["Very clear", "Mostly clear", "Somewhat unclear", "Unclear"])
        q3 = st.radio("Was the layout simple to follow?",
                      ["Yes", "Mostly", "No"])
        q4 = st.radio("Would you pay for this report?",
                      ["Yes", "Possibly", "No", "Already would have paid"])
        q5 = st.text_input("What would a fair price be for the full report? (optional)", placeholder="e.g. £15, £25")
        q6 = st.radio("Would you subscribe for ongoing access?",
                      ["Yes", "Possibly", "No"])
        q7 = st.radio("What personalisation matters most to you?",
                      ["Colour theme", "Text size", "Both", "Not bothered"])
        q8 = st.text_area("Anything else you'd like to tell us? (optional)", height=80)

        st.markdown("---")
        st.markdown("**Would you like to be notified when FRED launches?**")
        notify = st.radio("", ["Yes, notify me", "No thank you"], key="survey_notify")
        notify_email = ""
        if notify == "Yes, notify me":
            notify_email = st.text_input("Email address for launch notification", key="survey_email")

        submitted = st.form_submit_button("Submit feedback")
        if submitted:
            st.session_state.survey_submitted = True
            st.success("Thank you. Your feedback helps us build FRED properly.")


# ── CORRESPONDENCE MODULE ─────────────────────────────────────────────────────

def detect_tone_recommendation(text, patterns):
    red_count   = sum(1 for p in patterns if p["tier"] == "red")
    amber_count = sum(1 for p in patterns if p["tier"] == "amber")
    green_count = sum(1 for p in patterns if p["tier"] == "green")

    # ── Existing formal signals ───────────────────────────────────────────────
    formal_signals = re.compile(
        r"\b(following legal advice|our solicitor|legal services|"
        r"we are unable|complaints procedure|without prejudice|"
        r"we must inform|we are not in a position|governing body|"
        r"we refute|we dispute|we deny|permanent exclusion|"
        r"fixed term exclusion|resources do not allow)\b",
        re.IGNORECASE
    )

    # ── Extended formal signals — escalation and process substitution ─────────
    escalation_signals = re.compile(
        r"\b(escalated to|referred to educational psychology|"
        r"statutory assessment request|alternative provision referral|"
        r"behaviour support plan review|emergency annual review|"
        r"consultation period triggered|formal panel review|"
        r"multi.agency referral|risk assessment updated|"
        r"complex needs panel|managed move|high.needs funding application|"
        r"fixed.term suspension|safeguarding threshold|"
        r"retrospective|post.incident review|"
        r"records indicate prior patterns|evidence gathered over time|"
        r"within available resources|resource.limited intervention|"
        r"capacity.linked support|delivered within standard ratios|"
        r"shared teaching assistant|pro.rata provision delivery)\b",
        re.IGNORECASE
    )

    # ── Existing collaborative signals ────────────────────────────────────────
    collab_signals = re.compile(
        r"\b(we would like to|working together|we appreciate|"
        r"we understand your concerns|happy to discuss|"
        r"we are committed|please let us know|we value|"
        r"thank you for|we recognise|we are sorry|"
        r"we will look into|we want to support)\b",
        re.IGNORECASE
    )

    # ── Operational SEND language — school is engaging correctly ─────────────
    operational_signals = re.compile(
        r"\b(sensory diet|trusted adult|safe space|take.up time|"
        r"now and next|visual timetable|social stor(?:y|ies)|"
        r"non.confrontational|regulation break|emotion coaching|"
        r"co.regulation|zones of regulation|sensory circuit|"
        r"movement break|de.escalation|low arousal|check.in)\b",
        re.IGNORECASE
    )

    formal_count      = len(formal_signals.findall(text))
    escalation_count  = len(escalation_signals.findall(text))
    collab_count      = len(collab_signals.findall(text))
    operational_count = len(operational_signals.findall(text))

    # ── Scoring logic ─────────────────────────────────────────────────────────
    # Escalation language always pushes toward formal regardless of warm tone.
    # Collaborative mask: warm words + escalation = formal, not collaborative.
    # Operational language: school using correct SEND terms = slight collab signal.

    adjusted_formal = formal_count + (escalation_count * 2)
    adjusted_collab = collab_count + operational_count - (escalation_count * 1)

    # Collaborative mask detection — warm language paired with escalation
    collab_mask_active = (
        collab_count >= 2 and escalation_count >= 1
    )

    if collab_mask_active:
        return {
            "recommendation": "formal",
            "label": "Formal written response — collaborative mask detected",
            "reasoning": (
                "This correspondence uses warm, collaborative language alongside "
                "escalation or process language. This combination — partnership tone "
                "paired with referrals, panels, or resource restrictions — is a "
                "recognised pattern. The warm tone does not change the substance. "
                "A formal written response that focuses on provision specifics "
                "rather than the relationship is appropriate here. "
                "Put everything in writing and request written responses within "
                "five working days."
            ),
            "confidence": "high",
        }

    if red_count >= 2 or adjusted_formal >= 2:
        return {
            "recommendation": "formal",
            "label": "Formal written response",
            "reasoning": (
                f"The correspondence contains "
                f"{'serious provision or escalation patterns' if red_count >= 2 else 'formal or legal language'} "
                f"that suggest the school is managing rather than engaging. "
                f"A formal written response that references Section F specifically "
                f"is likely to be more effective than a collaborative one. "
                f"Put everything in writing and request a written response "
                f"within five working days."
            ),
            "confidence": "high" if red_count >= 2 and adjusted_formal >= 1 else "moderate",
        }

    elif adjusted_collab >= 2 and red_count == 0 and escalation_count == 0:
        # Only genuinely collaborative if no escalation present
        operational_note = (
            " The school is also using correct SEND language which is a positive signal."
            if operational_count >= 2 else ""
        )
        return {
            "recommendation": "collaborative",
            "label": "Collaborative response",
            "reasoning": (
                "The language in this correspondence is constructive and appears to be "
                "engaging with your concerns. A collaborative tone is likely to preserve "
                "the relationship while still holding the school to account."
                f"{operational_note} "
                "You can be warm and firm at the same time."
            ),
            "confidence": "moderate",
        }

    else:
        operational_note = (
            " The school is using correct SEND language — that is a positive sign "
            "but check it is backed by a delivery log."
            if operational_count >= 1 else ""
        )
        return {
            "recommendation": "neutral",
            "label": "Measured response",
            "reasoning": (
                "The correspondence does not show strong signals in either direction. "
                "A measured written response — factual, specific, referencing the EHCP — "
                f"is appropriate.{operational_note} "
                "Neither overly formal nor overly warm."
            ),
            "confidence": "low",
        }
def add_to_thread(direction, summary, patterns, tone_rec, reply_sent=""):
    if "thread" not in st.session_state:
        st.session_state.thread = []
    entry = {
        "date": datetime.datetime.now().strftime("%d %B %Y"),
        "direction": direction,
        "summary": summary,
        "patterns": [p["name"] for p in patterns],
        "tone_rec": tone_rec.get("recommendation", ""),
        "reply_sent": reply_sent,
    }
    st.session_state.thread.append(entry)
    summaries = [f"{e['date']}: {e['summary']}" for e in st.session_state.thread[-5:]]
    st.session_state.thread_context = " | ".join(summaries)


def check_thread_for_similar(patterns, environment):
    if not st.session_state.get("thread"):
        return []
    current_pattern_names = {p["name"] for p in patterns}
    similar = []
    for entry in st.session_state.thread[:-1]:
        past_names = set(entry.get("patterns", []))
        overlap = current_pattern_names & past_names
        if overlap or (environment and environment in entry.get("summary", "")):
            similar.append({**entry, "overlap": list(overlap)})
    return similar


# ── CORRESPONDENCE PATTERN LIBRARY ───────────────────────────────────────────

CORRESPONDENCE_PATTERNS = [
    {
        "id": "resources_defence",
        "name": "The Resources Defence",
        "tier": "red",
        "triggers": r"\b(resources|cost involved|beyond its requirements|school budget|cannot fund|not funded|staffing constraints|financial)\b",
        "explanation": "The school is suggesting that cost or resources limit what they can provide. The duty to deliver specified EHCP provision is absolute under the Children and Families Act 2014. It does not depend on the school's budget. Where additional funding is needed, the responsibility lies with the LA.",
        "fred_question": "Ask the school in writing to identify which provisions in Section F they are delivering and which they are not delivering due to resource constraints. Any gap becomes a matter for the LA to fund.",
    },
    {
        "id": "relationship_threat",
        "name": "The Relationship Breakdown Threat",
        "tier": "red",
        "triggers": r"\b(placement.*fail|break.*down|irrevocably|if.*situation.*not.*change|not.*engage|refuse.*communicate|consequences)\b",
        "explanation": "The school is implying that asserting your rights will damage the placement. This inverts responsibility. The duty to deliver provision exists regardless of whether parents are in dispute. A threat of placement breakdown in response to a parent exercising their rights is pressure, not professional dialogue.",
        "fred_question": "Note the date and exact wording. Respond in writing acknowledging receipt only. Do not respond to the threat. Ask instead for written confirmation of what provision is currently being delivered and what the support plan is.",
    },
    {
        "id": "best_interests_redirect",
        "name": "The Best Interests Redirect",
        "tier": "amber",
        "triggers": r"\b(best interest|not right for|another.*placement|more.*suitable|alternative.*provision|better.*placed|different.*school)\b",
        "explanation": "The school is using welfare language to suggest your child would be better placed elsewhere. Any professional opinion about placement must be made through the EHCP review process with evidence — not in correspondence responding to a complaint.",
        "fred_question": "Ask what evidence base supports this view, who assessed it, and whether it has been formally submitted to the LA as part of the EHCP review process.",
    },
    {
        "id": "blip_normalisation",
        "name": "The Blip Normalisation",
        "tier": "amber",
        "triggers": r"\b(only a blip|usual for his age|usual for her age|all children|typical.*behaviour|normal.*development|peers.*same|nothing.*unusual)\b",
        "explanation": "The school is comparing your child to neurotypical peers to minimise the significance of an incident. Your child has an EHCP precisely because their needs are not the same as their peers. 'Usual for their age' is not a relevant benchmark.",
        "fred_question": "Ask the school to confirm whether this assessment was made with reference to your child's EHCP profile and diagnosis. Ask them to reconsider in light of Section B of the plan.",
    },
    {
        "id": "complaint_deflection",
        "name": "The Complaint Policy Deflection",
        "tier": "amber",
        "triggers": r"\b(complaints policy|complaints procedure|formal complaint|pursue.*further|following.*legal.*advice|refer you to)\b",
        "explanation": "The school is directing you to the complaints procedure rather than engaging with the substance. The complaints procedure does not pause statutory obligations or EHCP timelines. It is a separate process.",
        "fred_question": "Note that the complaints procedure and the statutory EHCP process are separate. Continue through the EHCP route in parallel. Respond confirming you have noted the reference and are continuing to seek resolution of the provision concern.",
    },
    {
        "id": "good_week_signal",
        "name": "The Good Week Signal",
        "tier": "amber",
        "triggers": r"\b(good week|positive week|doing well|made good progress|responding well|settled week|happy in school)\b",
        "explanation": "Positive weekly summaries create a paper trail that can contradict the severity of your child's needs if an incident occurs in the same period. A log that records only positives is not an accurate delivery record.",
        "fred_question": "Ask for the full incident record alongside positive weekly summaries. If both a positive summary and a behavioural incident occurred in the same week, ask the school to explain the discrepancy.",
    },
    {
        "id": "reintegration_promise",
        "name": "The Reintegration Promise Without a Plan",
        "tier": "red",
        "triggers": r"\b(reintegration|return.*plan|back.*school|phased.*return|managed.*return|support.*return)\b",
        "explanation": "A verbal or vague reintegration promise is not a reintegration plan. A child with an EHCP who has been excluded requires a documented reintegration plan that references the EHCP provision. Without it there is no Do stage.",
        "fred_question": "Request the written reintegration plan within five working days, referencing the EHCP explicitly. If no written plan exists, confirm that in writing and ask when one will be provided.",
    },
    {
        "id": "staffing_change",
        "name": "Unstated Staffing Change",
        "tier": "red",
        "triggers": r"\b(new.*staff|change.*support|different.*adult|new.*TA|new.*teaching assistant|cover|temporary|supply)\b",
        "explanation": "For a child with ASD or similar needs, key adult relationships are often specified in the EHCP. A staffing change affecting 1:1 support is a provision change. It should not happen without parents being informed and without a transition plan.",
        "fred_question": "Ask the school to confirm in writing whether any staffing changes affecting your child's 1:1 support have been made since the last annual review, and what transition planning was put in place.",
    },
    {
        "id": "legal_misrepresentation",
        "name": "Legal Obligation Misrepresented",
        "tier": "red",
        "triggers": r"\b(no legal obligation|not required by law|not a legal requirement|beyond what is required|not legally obliged|discretionary)\b",
        "explanation": "The school may be correctly stating there is no obligation to provide 'the best possible education' — but this is not the relevant obligation. The obligation under the Children and Families Act 2014 is to deliver every provision specified in Section F. That obligation is absolute, not discretionary.",
        "fred_question": "Ask the school to confirm delivery of each named provision in Section F. The question is not whether they are providing the best possible education — it is whether they are delivering what the EHCP specifies.",
    },
    {
        "id": "homeschool_discrepancy",
        "name": "Home/School Discrepancy Used Against Parents",
        "tier": "amber",
        "triggers": r"\b(home.*school.*different|parents.*expectations|at home.*different|parental.*concern|parents.*report|home.*school.*gap)\b",
        "explanation": "The school is using differences between home and school presentation to question parental credibility. A child who presents differently at home and at school may be experiencing a stress response that is discharged at home. This is a clinical question, not a credibility question.",
        "fred_question": "Ask the school what steps they have taken to investigate the home/school discrepancy as a clinical question. Ask whether it has been formally assessed and what the findings were.",
    },
    {
        "id": "reassurance_without_evidence",
        "name": "Reassurance Without Evidence",
        "tier": "amber",
        "triggers": r"\b(we are not concerned|no concerns at this time|fully supported|meeting his needs|meeting her needs|you would be our first|runs continuously|we monitor closely)\b",
        "explanation": "The correspondence contains reassurance language without documentary support. Reassurance is not evidence. An assertion that provision is in place or that needs are being met requires a delivery log, not a statement of confidence.",
        "fred_question": "Ask what the evidence base is for this statement. Request the delivery log for the relevant provision over the past half term. If no log exists, note that in writing.",
    },
    {
        "id": "implicit_admission",
        "name": "Implicit Admission of Current Gap",
        "tier": "amber",
        "triggers": r"\b(going forward|in future|we will ensure|we have identified|we are working to improve|we recognise|we are reviewing|steps have been taken)\b",
        "explanation": "This language implies current practice is inadequate while presenting future improvement as reassurance. A commitment to improve in future is an admission of a current gap. Note the date — this correspondence is evidence that the gap existed at this point.",
        "fred_question": "Ask the school to confirm in writing what the current position is — not what it will be. Note that this correspondence has been retained and dated.",
    },
    {
        "id": "behaviour_framing",
        "name": "SEND Need Framed as Behaviour Problem",
        "tier": "red",
        "triggers": r"\b(behaviour.*pattern|pattern.*behaviour|conduct|expectations.*behaviour|"
                    r"struggling to meet expectations|behaviour.*concern|behavioural incident|"
                    r"standards expected|reminded.*standards|reflect on his behaviour|"
                    r"reflect on her behaviour|out of circulation|time to reflect)\b",
        "explanation": (
            "The school is describing what may be a SEND need as a behaviour problem. "
            "For a child with an EHCP, any incident during unstructured time — lunchtime, "
            "break, transitions — must first be examined against what provision was in place "
            "at that moment. A child with documented sensory or social communication needs "
            "who becomes dysregulated in an unstructured environment is not misbehaving — "
            "they are experiencing a provision failure."
        ),
        "fred_question": (
            "Ask the school what provision from Section F was in place during the unstructured "
            "period when this incident occurred. If 1:1 support, sensory breaks, or structured "
            "lunchtime activity is specified in the EHCP, ask whether it was being delivered "
            "at the time of the incident."
        ),
    },
    {
        "id": "veiled_threat",
        "name": "Veiled Threat — Next Steps Language",
        "tier": "red",
        "triggers": r"\b(next steps|may need to consider|if.*persist|if.*continue|"
                    r"cannot continue|placement.*appropriate|alternative.*provision|"
                    r"exploring.*options|reviewing.*placement|not the right environment|"
                    r"may not be meeting|current environment.*not)\b",
        "explanation": (
            "The correspondence contains language that implies escalation — toward exclusion, "
            "alternative provision, or placement change — dressed as concern. "
            "'Next steps in terms of support and provision' following a behavioural incident "
            "is a soft version of a placement threat. Note the date and exact wording. "
            "This creates a paper trail that the school may use to justify future action."
        ),
        "fred_question": (
            "Ask the school to clarify in writing what 'next steps' means specifically. "
            "Is the school considering a change to provision, a change to placement, or "
            "a referral? Each has different statutory implications. "
            "Any change to provision must go through the EHCP review process."
        ),
    },
    {
        "id": "home_responsibility_redirect",
        "name": "Home Responsibility Redirect",
        "tier": "amber",
        "triggers": r"\b(reinforce at home|support at home|work with.*at home|"
                    r"we would ask.*home|parental.*support|home.*behaviour|"
                    r"consistent.*home|same.*home|boundaries at home|routines at home)\b",
        "explanation": (
            "The school is redirecting responsibility for a school-based incident to the home. "
            "School behaviour during school hours is a school provision question, not a home "
            "behaviour question. An incident in the dining hall is governed by what provision "
            "was in place in the dining hall — not by what happens at home."
        ),
        "fred_question": (
            "Note this redirect. Respond by asking what provision was in place at school "
            "during the incident rather than engaging with the home behaviour framing. "
            "The EHCP provision is a school responsibility, not a shared one."
        ),
    },
    {
        "id": "unstructured_time_gap",
        "name": "Incident During Unstructured Time — Provision Check",
        "tier": "red",
        "triggers": r"\b(lunchtime|lunch time|break time|breaktime|unstructured|"
                    r"less structured|playground|dinner time|free time|between lessons|"
                    r"transition.*incident|incident.*transition|corridor.*incident)\b",
        "explanation": (
            "The incident occurred during an unstructured period — lunchtime, break, or "
            "transition. For a child with ASD, SEMH needs, or sensory processing difficulties, "
            "unstructured time is consistently the highest-risk period of the school day. "
            "The EHCP should specify what support is in place during these periods. "
            "If it does not, that is a provision gap. If it does and the support was not "
            "in place, that is a delivery failure."
        ),
        "fred_question": (
            "Ask the school to confirm in writing what provision from Section F is specifically "
            "in place during lunchtime, break time, and transitions. "
            "Then ask whether that provision was being delivered at the time of the incident. "
            "Request the delivery log for the relevant period."
        ),
    },
    {
        "id": "monitoring_without_action",
        "name": "Monitor and Observe — No Specified Action",
        "tier": "amber",
        "triggers": r"\b(monitor.*closely|continue to monitor|keep an eye|observe.*situation|"
                    r"watching.*closely|review.*situation|situation.*closely|monitor.*progress)\b",
        "explanation": (
            "The school has committed to monitoring without specifying what action will be "
            "taken or what threshold will trigger a review. 'We will monitor closely' is "
            "not a provision — it is an observation. Monitoring without a specified response "
            "plan is not a graduated approach under the SEND Code of Practice."
        ),
        "fred_question": (
            "Ask the school what specifically they will be monitoring, what the review "
            "mechanism is, who is responsible, and what action will be taken if the "
            "monitoring identifies further difficulty. Ask for this in writing."
        ),
    },
   {
        "id": "process_substitution",
        "name": "The Process Substitution",
        "tier": "red",
        "triggers": (
            r"\b(graduated response|universal provision|quality first teaching|"
            r"whole.school approach|standard tracking|differentiation within|"
            r"pastoral support plan|wave 1 intervention|internal monitoring|"
            r"delivered within standard timetable|curriculum modification|"
            r"monitored by the class teacher|flexible grouping|"
            r"scaffolding strategies|in-house literacy screening)\b"
        ),
        "explanation": (
            "The school has shifted from SEND provision language to operational and "
            "policy management language. Terms like 'graduated response', 'quality first "
            "teaching', and 'universal provision' describe what the school does for all "
            "pupils — they are not the specified, individualised provision required by "
            "an EHCP. A child with an EHCP is entitled to provision above and beyond "
            "universal classroom practice. Substituting process language for provision "
            "language is a way of appearing to respond without committing to anything "
            "that is specific, quantified, or enforceable."
        ),
        "fred_question": (
            "Ask the school to identify which provisions in Section F of the EHCP — "
            "by name, frequency, and deliverer — are currently being delivered. "
            "The question is not what the school does generally. "
            "It is what is being delivered for this child, as specified in the plan."
        ),
    },
    {
        "id": "escalation_sequence",
        "name": "The Escalation Sequence",
        "tier": "red",
        "triggers": (
            r"(?:sensory diet|trusted adult|safe space|take[\s-]up time|"
            r"now and next|now-and-next|visual timetable|social stor(?:y|ies)|"
            r"non[\s-]confrontational|regulation break|emotion coaching|"
            r"co[\s-]regulation|zones of regulation|sensory circuit|"
            r"movement break|de[\s-]escalation|low arousal|check[\s-]in|"
            r"safe adult|exit card|emotion thermometer|5[\s-]point scale|"
            r"feelings fan|time out card|anxiety support|emotional regulation)"
        ),
        "explanation": (
            "The correspondence contains escalation terminology — referrals, panels, "
            "risk assessments, managed moves. These are formal process steps that "
            "have significant consequences for your child's placement and provision. "
            "Escalation language in correspondence often appears before a formal "
            "decision has been made, creating a paper trail that justifies later action. "
            "Each escalation step has a statutory process — it cannot happen without "
            "proper notice, evidence, and parental involvement."
        ),
        "fred_question": (
            "Ask the school or LA to confirm in writing what process is being initiated, "
            "what the statutory basis for it is, and what your rights are at each stage. "
            "Ask for copies of all referral documents and risk assessments. "
            "Any change to provision or placement must go through the EHCP review process "
            "under the Children and Families Act 2014."
        ),
    },
    {
        "id": "collaborative_mask",
        "name": "The Collaborative Mask",
        "tier": "amber",
        "triggers": (
            r"\b(working collaboratively|partnership approach|co.production|"
            r"working in tandem|shared expectations|collaborative solution|"
            r"team around the child|open dialogue|working towards common goals|"
            r"parental views incorporated|joint action planning|"
            r"transparent communication|multi.disciplinary approach|"
            r"liaising closely|holistic support model|agreed next steps|"
            r"shared responsibility|consensus building|child.centred review)\b"
        ),
        "explanation": (
            "The correspondence uses warm collaborative language — partnership, "
            "co-production, open dialogue, shared goals. This tone is not itself a "
            "problem. The concern is when collaborative language appears alongside "
            "procedural distancing, escalation steps, or provision substitution. "
            "Warm words do not pause statutory obligations. A letter that opens with "
            "'we are committed to working in partnership' and closes with a referral "
            "to a panel is not a collaborative document — it is a managed one."
        ),
        "fred_question": (
            "Read the full correspondence for what is being proposed or initiated, "
            "not just the tone. Ask the school to confirm specifically what will "
            "change in terms of provision — not the relationship, the provision. "
            "Collaborative language without a specific provision commitment "
            "is reassurance, not a plan."
        ),
    },
    {
        "id": "retrospective_justification",
        "name": "The Retrospective Justification",
        "tier": "red",
        "triggers": (
            r"\b(retrospective|post.incident review|following a period of monitoring|"
            r"subsequent data analysis|records indicate prior patterns|"
            r"historical tracking|in light of recent assessments|"
            r"evidence gathered over time|retrospective application|"
            r"incident log reconciliation|following review of previous strategies|"
            r"patterns established through data|chronological review|"
            r"subsequent professional advice|post.hoc evaluation|"
            r"re.evaluation of previous entries|file review demonstrates|"
            r"longitudinal tracking confirms|evidence base established post.event)\b"
        ),
        "explanation": (
            "The correspondence contains retrospective justification language — "
            "the school is presenting a later reconstruction of events, provision, "
            "or process as if it were contemporaneous evidence. Phrases like "
            "'records indicate prior patterns' or 'evidence gathered over time reveals' "
            "suggest that documentation is being assembled after the fact. "
            "Contemporaneous records — delivery logs, incident reports, session notes — "
            "written at the time carry significantly more evidential weight than "
            "retrospective summaries. If a pattern was identified over time, "
            "ask when it was first recorded and what action was taken at that point."
        ),
        "fred_question": (
            "Ask the school to provide the original contemporaneous records — "
            "the delivery logs, session notes, and incident reports written at the time, "
            "not a summary compiled afterwards. Ask specifically: when was this pattern "
            "first recorded, who recorded it, and what action was taken at that point? "
            "If no contemporaneous record exists, note that in writing."
        ),
    },
    {
        "id": "provision_substitution_corr",
        "name": "The Provision Substitution",
        "tier": "red",
        "triggers": (
            r"\b(within available resources|flexible delivery model|"
            r"equivalent alternative provision|integrated support delivery|"
            r"delivered within standard ratios|shared teaching assistant|"
            r"group.based provision model|delivered within existing frameworks|"
            r"alternative timetabling|resource allocation panel|"
            r"delivered via general school environment|"
            r"pro.rata provision delivery|integrated therapeutic model|"
            r"delivered by available pastoral staff|flexible group structures|"
            r"capacity.linked support|resource.matched intervention|"
            r"delivered within current site constraints|"
            r"alternative curriculum pathway|delivered via class.based deployment|"
            r"integrated within standard delivery|resource.limited intervention)\b"
        ),
        "explanation": (
            "The school is describing provision in terms of what is operationally "
            "convenient rather than what the EHCP specifies. Phrases like 'within "
            "available resources', 'shared teaching assistant support', and 'flexible "
            "delivery model' are substitutes for the named, quantified, specific "
            "provision required by Section F. The duty to deliver is absolute under "
            "the Children and Families Act 2014 — it is not modifiable by the school's "
            "staffing model, timetable, or resource allocation decisions. "
            "If the school cannot deliver what Section F specifies within its current "
            "structure, the obligation passes to the LA to make alternative arrangements."
        ),
        "fred_question": (
            "Ask the school to confirm which provisions in Section F are being delivered "
            "exactly as specified — named role, stated frequency, stated duration. "
            "For any provision described as 'flexible', 'shared', or 'within available "
            "resources', ask what the Section F specification actually says and how "
            "that differs from what is currently being delivered."
        ),
    },
    {
        "id": "operational_language_signal",
        "name": "Correct SEND Language Used",
        "tier": "green",
        "triggers": (
            r"\b(sensory diet|trusted adult|safe space|take.up time|"
            r"now and next|visual timetable|social stor(?:y|ies)|"
            r"non.confrontational|regulation break|emotion coaching|"
            r"co.regulation|zones of regulation|sensory circuit|"
            r"movement break|de.escalation|low arousal|check.in)\b"
        ),
        "explanation": (
            "The school is using specific SEND provision language in this correspondence. "
            "This is what good looks like — it suggests staff are familiar with your "
            "child's needs and are using the right framework to describe support. "
            "This is a positive signal and should be acknowledged."
        ),
        "fred_question": (
            "The delivery check: this language should appear in Section F of the EHCP "
            "as a specified, quantified commitment — not just in correspondence. "
            "Ask the school to confirm which of these provisions are formally written "
            "into the plan, how often they are delivered, by whom, and whether a "
            "delivery log is being kept. Language in an email is not the same as "
            "a lawful commitment in the EHCP. Good language plus a delivery log "
            "is the standard. Good language alone is a start."
        ),
    }, 
]

ENVIRONMENT_TRIGGERS = {
    "food hall": "food hall / canteen",
    "canteen": "food hall / canteen",
    "dining hall": "food hall / canteen",
    "dining room": "food hall / canteen",
    "lunch hall": "food hall / canteen",
    "corridor": "corridor / transition space",
    "hallway": "corridor / transition space",
    "playground": "playground / outdoor space",
    "outside": "playground / outdoor space",
    "assembly": "assembly hall",
    "hall": "assembly hall",
    "classroom": "classroom environment",
    "gym": "sports hall / gym",
    "sports hall": "sports hall / gym",
    "pe": "sports hall / gym",
    "toilet": "toilet facilities",
    "bathroom": "toilet facilities",
    "bus": "school transport",
    "minibus": "school transport",
}

SENSORY_CHECKLIST = {
    "food hall / canteen": [
        ("Sound", [
            "Overall noise level — crowded dining hall acoustics",
            "Cutlery, chair scraping, tray noise",
            "Echo and reverberation in the space",
            "Kitchen extraction fans and cooking sounds",
            "Unpredictable loud sounds — dropped trays, shouting",
        ]),
        ("Light", [
            "Strip lighting or fluorescent flicker",
            "Overall brightness",
            "Glare from windows or surfaces",
        ]),
        ("Smell", [
            "Food odours — specific foods that may be aversive",
            "Cleaning products",
            "Accumulated smell of multiple food types",
        ]),
        ("Physical", [
            "Crowding — proximity to other pupils",
            "Queue management — unpredictable movement",
            "Seating — choice or assigned, proximity to others",
            "Clear exit route visible",
        ]),
        ("Timing", [
            "Hunger state on arrival — time since last food",
            "Blood sugar — snack provided before lunch?",
            "Preceding activity — what happened immediately before",
        ]),
        ("Predictability", [
            "Does the child know what food is available before arriving",
            "Consistent routine day to day",
            "Visual support for the lunch process",
        ]),
    ],
    "corridor / transition space": [
        ("Sound", [
            "Noise level during transition — multiple classes moving",
            "Unpredictable sounds — lockers, doors, shouting",
            "Echo in narrow spaces",
        ]),
        ("Physical", [
            "Crowding during peak transition times",
            "Physical contact from other pupils",
            "Width of corridor — personal space available",
        ]),
        ("Predictability", [
            "Consistency of route taken",
            "Whether the child knows destination in advance",
            "Early departure to avoid peak crowding",
            "Named adult support during transition",
        ]),
    ],
    "playground / outdoor space": [
        ("Sound", [
            "Unpredictable noise — shouting, ball games",
            "Wind noise",
        ]),
        ("Social", [
            "Unstructured peer interaction — no adult mediation",
            "Safe space available if overwhelmed",
            "Structured activity available as alternative",
        ]),
        ("Sensory", [
            "Bright sunlight — visual sensitivity",
            "Temperature extremes",
            "Physical play — proprioceptive seeking or avoiding",
        ]),
    ],
}

DEFAULT_CHECKLIST = [
    ("Sound", ["Overall noise level", "Unpredictable sounds", "Echo or reverberation"]),
    ("Light", ["Brightness", "Fluorescent lighting", "Glare"]),
    ("Smell", ["Specific odours that may be aversive"]),
    ("Physical", ["Crowding", "Personal space", "Exit routes"]),
    ("Predictability", ["Routine consistency", "Advance warning of changes"]),
    ("Timing", ["Time of day", "Preceding activity", "Hunger or fatigue state"]),
]


def detect_environment(text):
    text_lower = text.lower()
    for trigger, env_name in ENVIRONMENT_TRIGGERS.items():
        if trigger in text_lower:
            return env_name
    return None


def get_checklist(env_name):
    return SENSORY_CHECKLIST.get(env_name, DEFAULT_CHECKLIST)


def detect_patterns(text, ehcp_text=""):
    matched = []
    for pattern in CORRESPONDENCE_PATTERNS:
        if re.search(pattern["triggers"], text, re.IGNORECASE):
            m = re.search(pattern["triggers"], text, re.IGNORECASE)
            ctx = ""
            if m:
                ctx = text[max(0, m.start()-100):m.end()+120].strip()
            matched.append({**pattern, "extract": ctx[:250]})

    if ehcp_text and re.search(
        r"\b(lunchtime|lunch time|break time|unstructured|dining hall|canteen|playground)\b",
        text, re.IGNORECASE
    ):
        ehcp_lunch = re.search(
            r"\b(lunchtime|unstructured|break.*supervision|lunch.*support|"
            r"lunch.*adult|supervision.*lunch|1:1.*lunch|full.time.*support)\b",
            ehcp_text, re.IGNORECASE
        )
        if ehcp_lunch:
            matched.append({
                "id": "ehcp_xref_lunch",
                "name": "EHCP Cross-Reference — Lunchtime Provision",
                "tier": "red",
                "triggers": "",
                "explanation": (
                    "Your EHCP references provision during lunchtime or unstructured periods. "
                    "This incident occurred during exactly that period. "
                    "The question is not whether the behaviour occurred — it is whether the "
                    "provision specified in the EHCP was being delivered when it did."
                ),
                "fred_question": (
                    "Request the delivery log for lunchtime provision on the day of this incident. "
                    "Ask the school to confirm in writing whether the named provision from "
                    "Section F was in place at the time. If it was not, that is a delivery "
                    "failure, not a behaviour problem."
                ),
                "extract": "",
            })
        else:
            matched.append({
                "id": "ehcp_xref_lunch_gap",
                "name": "EHCP Gap — No Lunchtime Provision Specified",
                "tier": "red",
                "triggers": "",
                "explanation": (
                    "This incident occurred during lunchtime but the EHCP does not appear "
                    "to specify provision for this period. For a child with documented "
                    "sensory or social communication needs, unstructured lunchtime is a "
                    "high-risk period. The absence of specified provision is itself a gap "
                    "that should be addressed at the next annual review."
                ),
                "fred_question": (
                    "At the next annual review, request that Section F is amended to specify "
                    "what support is in place during all unstructured periods — lunchtime, "
                    "break, and transitions. Ask the school in writing what support was "
                    "in place at the time of this incident."
                ),
                "extract": "",
            })

    order = {"red": 0, "amber": 1, "green": 2}
    matched.sort(key=lambda x: order.get(x["tier"], 3))
    return matched


def analyse_policy(policy_text):
    findings = []
    audit_pattern = re.compile(
        r"\b(will conduct|will undertake|will carry out|will review|annual.*audit|"
        r"acoustic.*survey|accessibility.*audit|risk.*assessment.*will|will.*assess|plan to|intend to)\b",
        re.IGNORECASE
    )
    for m in audit_pattern.finditer(policy_text):
        ctx = policy_text[max(0, m.start()-60):m.end()+100].strip()
        findings.append({
            "tier": "amber",
            "title": "Policy commitment — check whether actioned",
            "extract": ctx[:280],
            "commentary": (
                "This policy commits to an action — audit, survey, or review. "
                "Ask the school for evidence it has been completed and the date it was last carried out. "
                "If it has not been done, this is a gap between policy commitment and practice."
            ),
        })
        break

    year_pattern = re.compile(r"\b(20\d{2})\b")
    years = [int(y) for y in year_pattern.findall(policy_text)]
    current_year = datetime.datetime.now().year
    if years:
        most_recent = max(years)
        age = current_year - most_recent
        if age >= 2:
            findings.append({
                "tier": "amber",
                "title": f"Policy is {age} years old — commitments may be overdue",
                "extract": f"Most recent year reference: {most_recent}",
                "commentary": (
                    f"This policy was last updated in {most_recent}. Any commitments — "
                    f"audits, assessments, training — should have been actioned in the "
                    f"intervening {age} years. Ask for evidence of actions taken since that date."
                ),
            })

    if re.search(r"\b(acoustic|strip light|fluorescent|sensory|ear defender|quiet area|calm space|accessibility)\b", policy_text, re.IGNORECASE):
        findings.append({
            "tier": "green",
            "title": "Policy references sensory or accessibility provision",
            "extract": "",
            "commentary": (
                "The policy references sensory or accessibility provision. "
                "Cross-reference this against your child's EHCP Section F. "
                "If the policy commits to provision that is absent from Section F, "
                "that gap is evidenced by the school's own document."
            ),
        })

    if re.search(r"\b(staff.*training|trained.*staff|autism.*training|SEND.*training|awareness.*training)\b", policy_text, re.IGNORECASE):
        findings.append({
            "tier": "amber",
            "title": "Policy commits to staff training — ask for evidence",
            "extract": "",
            "commentary": (
                "The policy references staff training in relation to SEND or specific needs. "
                "Ask the school to confirm what training has been provided to staff working "
                "directly with your child, and when it was last updated."
            ),
        })

    return findings


def render_pattern_card(p, index):
    tier = p["tier"]
    tier_label = (
        "RED — Lawful concern" if tier == "red"
        else "GREEN — Positive signal" if tier == "green"
        else "AMBER — Pattern detected"
    )
    st.markdown(f"""
    <div class="finding-{tier}" style="margin-bottom:1.2rem;">
      <span class="badge-{tier}">{tier_label}</span>
      <p style="font-weight:700;margin:0.5rem 0 0.3rem;font-size:1rem;">{p['name']}</p>
      <p style="margin:0 0 0.6rem;font-size:0.93rem;line-height:1.6;">{p['explanation']}</p>
      <div style="background:rgba(0,0,0,0.04);border-radius:4px;padding:0.6rem 0.8rem;margin-bottom:0.4rem;">
        <p style="font-weight:700;margin:0 0 0.2rem;font-size:0.88rem;text-transform:uppercase;letter-spacing:0.05em;">The question to ask:</p>
        <p style="margin:0;font-size:0.93rem;font-style:italic;">{p['fred_question']}</p>
      </div>
      {f'<p style="font-size:0.82rem;color: #666;margin:0.4rem 0 0;border-left:3px solid #ddd;padding-left:0.6rem;">{p["extract"]}…</p>' if p.get("extract") else ""}
    </div>
    """, unsafe_allow_html=True)


def generate_amendment_word(environment, confirmed_items, today):
    doc = Document()
    title = doc.add_heading("EHCP Amendment Request", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {today}")
    doc.add_paragraph(f"Re: Request to amend Section B and Section F — {environment}")
    doc.add_paragraph("")

    doc.add_paragraph("Dear [SENCO name],")
    doc.add_paragraph("")
    doc.add_paragraph(
        f"I am writing to request a formal amendment to my child's Education, Health and Care Plan, "
        f"specifically to Sections B and F, following observations made in the {environment} environment."
    )
    doc.add_paragraph("")
    doc.add_paragraph("The following factors have been identified through observation:")

    for item in confirmed_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph("")
    doc.add_paragraph(
        "These observations correspond to needs already described in Section B of the current plan. "
        "The current Section F does not specify provision for this environment in sufficient detail "
        "to be enforceable. I am requesting that the following be added to Section F:"
    )

    additions = doc.add_paragraph(style="List Bullet")
    additions.add_run(f"Specific provision for the {environment} environment, including named adjustments")
    doc.add_paragraph(style="List Bullet").add_run("Named arrangements to be in place on every occasion my child uses this space")
    doc.add_paragraph(style="List Bullet").add_run("A half-termly review mechanism to confirm adjustments are in place")

    doc.add_paragraph("")
    doc.add_paragraph(
        "I would be grateful for written confirmation that this request has been received "
        "and will be considered at the earliest opportunity."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph("[Your name]")
    doc.add_paragraph(f"[Date: {today}]")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def page_correspondence():
    """Correspondence intelligence page."""

    st.markdown("## Correspondence analysis")

    if st.session_state.get("findings"):
        if st.button("← Back to my EHCP report", key="back_to_report"):
            st.session_state.stage = "full_report"
            st.rerun()

    st.markdown(f"""
    <div style="background:#f0f4fa;border-radius:8px;padding:1.2rem 1.5rem;margin-bottom:1.5rem;">
      <p style="margin:0;font-size:0.97rem;line-height:1.7;">
        Upload a letter or email from school or the LA and FRED will read it for you —
        identifying patterns, root causes, and the right question to ask next.
        This is different from your EHCP report: the EHCP report analyses the document,
        correspondence analysis reads the relationship.
      </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### Upload correspondence")

    # REPAIR 3: key="input_method" now persists via session state default
    input_method = st.radio(
        "How do you want to add your correspondence?",
        ["Upload file (PDF or Word)", "Paste text directly"],
        horizontal=True,
        key="input_method",
    )

    col1, col2 = st.columns(2)

    # Initialise paste vars
    paste_text1 = ""
    paste_date1 = ""
    paste_text2 = ""
    paste_date2 = ""
    email1 = None
    email2 = None
    policy_file = None

    if input_method == "Upload file (PDF or Word)":
        with col1:
            email1 = st.file_uploader(
                "Most recent letter or email",
                type=["pdf", "docx", "txt"], key="email1"
            )
            email2 = st.file_uploader(
                "Earlier correspondence — optional",
                type=["pdf", "docx", "txt"], key="email2"
            )
        with col2:
            policy_file = st.file_uploader(
                "School policy — optional",
                type=["pdf", "docx", "txt"], key="policy_upload",
                help="Upload the school accessibility, SEN, or behaviour policy."
            )
            st.markdown(
                "<p style='font-size:0.82rem;color: #666;margin-top:0.3rem;'>"
                "Upload the school's accessibility or SEN policy and FRED will check whether "
                "they have met their own commitments."
                "</p>", unsafe_allow_html=True
            )
    else:
        with col1:
            st.markdown("**Most recent correspondence**")
            paste_date1 = st.text_input(
                "Date of this correspondence",
                placeholder="e.g. 12 May 2025",
                key="paste_date1"
            )
            paste_text1 = st.text_area(
                "Paste the text here",
                placeholder="Paste the full text of the email or letter...",
                height=180,
                key="paste_text1"
            )
        with col2:
            st.markdown("**Earlier correspondence — optional**")
            paste_date2 = st.text_input(
                "Date of this correspondence",
                placeholder="e.g. 3 April 2025",
                key="paste_date2"
            )
            paste_text2 = st.text_area(
                "Paste the text here",
                placeholder="Paste the full text of an earlier email or letter...",
                height=180,
                key="paste_text2"
            )

    st.markdown("---")

    tone_override = st.session_state.get("tone_override", None)
    if tone_override:
        st.markdown(f"""
        <div style="background:#eaf5e0;border-radius:8px;padding:0.6rem 1rem;margin-bottom:0.5rem;">
          <p style="margin:0;font-size:0.85rem;color: #2d5a2d;">
            Tone set to: <b>{tone_override}</b>
          </p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    analyse = st.button("Analyse correspondence", use_container_width=False, key="analyse_corr")

    # REPAIR 2: has_content checks both file upload and paste text
    has_content = (email1 is not None) or bool(paste_text1.strip())

    if not has_content and not analyse:
        st.markdown(
            "<p style='color:#888;font-size:0.9rem;margin-top:0.5rem;'>"
            "Upload at least one piece of correspondence, or paste the text above."
            "</p>",
            unsafe_allow_html=True
        )

    if analyse and not has_content:
        st.error("Please upload a file or paste some correspondence text to continue.")

    if analyse and has_content:
        st.session_state.corr_analysed = False

        progress_text = st.empty()
        progress_text.markdown("*Reading correspondence…*")

        if email1:
            text1 = extract_text(email1)
            if text1 == "__PDF_EXTRACTION_FAILED__":
                progress_text.empty()
                st.warning(
                    "FRED couldn't extract text from this PDF — it may be an image-based file. "
                    "Try using the 'Paste text directly' option instead, or save the email as a "
                    "Word document (.docx) and upload that."
                )
                st.stop()
        else:
            # REPAIR 2: use paste_text1 directly — it's already in scope
            date_prefix = f"[Date: {paste_date1}]\n" if paste_date1 else ""
            text1 = date_prefix + paste_text1

        if email2:
            text2 = extract_text(email2)
            if text2 == "__PDF_EXTRACTION_FAILED__":
                text2 = ""
        elif paste_text2.strip():
            date_prefix2 = f"[Date: {paste_date2}]\n" if paste_date2 else ""
            text2 = date_prefix2 + paste_text2
        else:
            text2 = ""

        # Store extracted text in session state before rerun
        st.session_state["paste_text_stored"] = text1 + "\n\n" + text2
        combined = st.session_state["paste_text_stored"]

        progress_text.markdown("*Identifying patterns…*")
        ehcp_text = st.session_state.get("vault", {}).get("ehcp", {}).get("text", "")
        matched_patterns = detect_patterns(combined, ehcp_text)

        progress_text.markdown("*Checking environment…*")
        environment = detect_environment(combined)

        policy_text = ""
        policy_findings = []
        if policy_file:
            progress_text.markdown("*Reading school policy…*")
            policy_text = extract_text_from_pdf(policy_file) if policy_file.name.endswith(".pdf") else extract_text_from_docx(policy_file)
            policy_findings = analyse_policy(policy_text)

        progress_text.markdown("*Detecting tone…*")
        tone_rec = detect_tone_recommendation(combined, matched_patterns)

        progress_text.markdown("*Checking correspondence history…*")
        similar_past = check_thread_for_similar(matched_patterns, environment)

        progress_text.empty()

        today = datetime.datetime.now().strftime("%d %B %Y")

        summary = f"Correspondence — {len(matched_patterns)} pattern(s) detected"
        if environment:
            summary += f" — {environment} referenced"
        existing = st.session_state.get("thread", [])
        if not existing or existing[-1].get("summary") != summary:
            add_to_thread("from_school", summary, matched_patterns, tone_rec)

        # Build draft reply
        red_n   = sum(1 for p in matched_patterns if p["tier"] == "red")
        amber_n = sum(1 for p in matched_patterns if p["tier"] == "amber")

        PRIORITY_IDS = [
            "behaviour_framing", "unstructured_time_gap", "ehcp_xref_lunch",
            "ehcp_xref_lunch_gap", "veiled_threat", "reintegration_promise",
            "staffing_change", "legal_misrepresentation", "resources_defence",
            "monitoring_without_action", "home_responsibility_redirect",
            "reassurance_without_evidence", "implicit_admission",
        ]

        QUESTION_PARAGRAPHS = {
            "behaviour_framing": (
                "I would be grateful if you could confirm in writing what provision from "
                "Section F of [child's name]'s EHCP was specifically in place at the time of "
                "this incident, and provide the delivery log for that period."
            ),
            "unstructured_time_gap": (
                "Could you confirm in writing what support was in place during lunchtime on "
                "the day of this incident, and share the relevant delivery records with me?"
            ),
            "ehcp_xref_lunch": (
                "The EHCP specifies provision during unstructured periods including lunchtime. "
                "I would ask you to confirm whether that provision was being delivered at the "
                "time of this incident, and to provide the delivery log."
            ),
            "ehcp_xref_lunch_gap": (
                "I note that the EHCP does not currently specify provision for lunchtime. "
                "Given that this incident occurred during that period, I would like to raise "
                "this at the next annual review as a gap that needs to be addressed."
            ),
            "veiled_threat": (
                "I would ask you to clarify in writing what 'next steps in terms of support "
                "and provision' means specifically. Any change to provision or placement must "
                "go through the EHCP review process under the Children and Families Act 2014."
            ),
            "reintegration_promise": (
                "Could you provide the written reintegration plan, specifically referencing "
                "the provision in Section F of the EHCP? I would ask for this within "
                "five working days."
            ),
            "staffing_change": (
                "Could you confirm whether there have been any changes to [child's name]'s "
                "1:1 support since the last annual review, and what transition planning "
                "was put in place for any such changes?"
            ),
            "legal_misrepresentation": (
                "I would ask the school to confirm delivery of each provision specified in "
                "Section F of the EHCP. The relevant obligation under the Children and "
                "Families Act 2014 is to deliver what the EHCP specifies."
            ),
            "resources_defence": (
                "Could you identify in writing which provisions in Section F are currently "
                "being delivered and whether any are not being delivered due to resource "
                "constraints? Any such gaps become a matter for the LA to address."
            ),
            "monitoring_without_action": (
                "Could you clarify what specifically will be monitored, what the review "
                "mechanism is, who is responsible, and what action will be taken if the "
                "monitoring identifies further difficulty? I would ask for this in writing."
            ),
            "home_responsibility_redirect": (
                "I would like to focus on what provision was in place at school at the time "
                "of the incident, as this is governed by the EHCP. Could you confirm this "
                "in writing?"
            ),
            "reassurance_without_evidence": (
                "Could you share the evidence base for the assurances given in your email, "
                "including the delivery log for the relevant provision over the past half term?"
            ),
            "implicit_admission": (
                "Could you confirm in writing what the current position is — specifically "
                "what is in place now, rather than what will be put in place going forward?"
            ),
        }

        tone_label = st.session_state.get("tone_override", tone_rec.get("recommendation", "neutral"))
        priority_sorted = sorted(
            [p for p in matched_patterns if p["id"] in PRIORITY_IDS],
            key=lambda x: PRIORITY_IDS.index(x["id"]) if x["id"] in PRIORITY_IDS else 99
        )
        top_two = priority_sorted[:2]

        openings = {
            "formal": "Dear [Name],\n\nThank you for your email dated [date]. I am writing to follow up on the points raised.",
            "collaborative": "Dear [Name],\n\nThank you for letting us know. I wanted to follow up on a couple of points.",
            "neutral": "Dear [Name],\n\nThank you for your email dated [date]. I would like to raise a couple of points in response.",
        }
        opening = openings.get(tone_label, openings["neutral"])
        intro = "I have a couple of specific questions I would be grateful if you could address in writing."
        question_lines = []
        for i, p in enumerate(top_two):
            q = QUESTION_PARAGRAPHS.get(p["id"], "")
            if q:
                question_lines.append(f"{i+1}. {q}")
        closings = {
            "formal": "I would be grateful for a written response within five working days.\n\nYours sincerely,\n[Your name]",
            "collaborative": "I look forward to hearing from you.\n\nKind regards,\n[Your name]",
            "neutral": "I look forward to your response.\n\nYours sincerely,\n[Your name]",
        }
        closing = closings.get(tone_label, closings["neutral"])

        if question_lines:
            draft_parts = [opening, intro] + question_lines + [closing]
        else:
            draft_parts = [opening, closing]
        draft_reply = "\n\n".join(p for p in draft_parts if p.strip())

        # Store everything in session state
        st.session_state["draft_reply"]        = draft_reply
        st.session_state["top_two_patterns"]   = top_two
        st.session_state["all_patterns"]       = matched_patterns
        st.session_state["corr_summary_bar"]   = (red_n, amber_n, "", "")
        st.session_state["corr_environment"]   = environment
        st.session_state["corr_policy"]        = policy_findings
        st.session_state["corr_tone_rec"]      = tone_rec
        st.session_state["corr_similar"]       = similar_past
        st.session_state["corr_analysed"]      = True
        st.rerun()

    # ── Display results ───────────────────────────────────────────────────────
    if st.session_state.get("corr_analysed"):
        matched_patterns = st.session_state.get("all_patterns", [])
        environment      = st.session_state.get("corr_environment")
        policy_findings  = st.session_state.get("corr_policy", [])
        tone_rec         = st.session_state.get("corr_tone_rec", {})
        draft_reply      = st.session_state.get("draft_reply", "")
        top_two          = st.session_state.get("top_two_patterns", [])
        today            = datetime.datetime.now().strftime("%d %B %Y")

        red_n   = sum(1 for p in matched_patterns if p["tier"] == "red")
        amber_n = sum(1 for p in matched_patterns if p["tier"] == "amber")

        green_n = sum(1 for p in matched_patterns if p["tier"] == "green")

        if red_n == 0 and amber_n == 0 and green_n == 0:
            summary_colour = GREEN
            summary_text = "No major patterns detected in this correspondence."
        elif red_n > 0:
            green_note = f" {green_n} positive signal{'s' if green_n != 1 else ''}." if green_n > 0 else ""
            summary_colour = RED
            summary_text = (
                f"{red_n} serious pattern{'s' if red_n != 1 else ''} and "
                f"{amber_n} amber signal{'s' if amber_n != 1 else ''} detected."
                f"{green_note}"
            )
        elif amber_n > 0:
            green_note = f" {green_n} positive signal{'s' if green_n != 1 else ''}." if green_n > 0 else ""
            summary_colour = AMBER
            summary_text = (
                f"{amber_n} pattern{'s' if amber_n != 1 else ''} detected. "
                f"No immediate lawful concerns."
                f"{green_note}"
            )
        else:
            summary_colour = GREEN
            summary_text = f"{green_n} positive signal{'s' if green_n != 1 else ''} detected. No concerns identified."

        st.markdown(f"""
        <div style="background:{summary_colour};border-radius:6px;padding:0.9rem 1.2rem;margin:1rem 0 1.5rem;">
          <p style="color:white;font-weight:700;margin:0;font-size:1rem;">{summary_text}</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("### Patterns detected")

        if not matched_patterns:
            st.markdown(f"""
            <div class="finding-green">
              <span class="badge-green">No recognised patterns</span>
              <p style="margin:0.5rem 0 0;font-size:0.95rem;">
                No recognised school correspondence patterns were detected.
                Read the correspondence carefully before assuming it is without issue.
              </p>
            </div>
            """, unsafe_allow_html=True)
        else:
            show_all = st.session_state.get("show_all_patterns", False)
            display = matched_patterns if show_all else matched_patterns[:5]

            for i, p in enumerate(display):
                render_pattern_card(p, i)

            if len(matched_patterns) > 5 and not show_all:
                remaining = len(matched_patterns) - 5
                if st.button(f"Show {remaining} more pattern{'s' if remaining > 1 else ''}"):
                    st.session_state.show_all_patterns = True
                    st.rerun()
            elif show_all and len(matched_patterns) > 5:
                if st.button("Show fewer"):
                    st.session_state.show_all_patterns = False
                    st.rerun()

        if environment:
            st.markdown("---")
            st.markdown(f"### Root cause — {environment}")
            st.markdown(f"""
            <div style="background:#f0f4fa;border-radius:6px;padding:0.9rem 1.2rem;margin-bottom:1rem;">
              <p style="margin:0;font-size:0.95rem;">
                The correspondence references <b>{environment}</b> as a context for difficulty.
                Tick everything that applies or has been observed.
              </p>
            </div>
            """, unsafe_allow_html=True)

            checklist = get_checklist(environment)
            confirmed_items = []

            for category, items in checklist:
                st.markdown(f"**{category}**")
                cols = st.columns(2)
                for idx, item in enumerate(items):
                    with cols[idx % 2]:
                        if st.checkbox(item, key=f"chk_{category}_{idx}"):
                            confirmed_items.append(f"{category}: {item}")

            if confirmed_items:
                st.markdown("---")
                st.markdown("### Root cause confirmed")

                st.markdown(f"""
                <div class="finding-red">
                  <span class="badge-red">{len(confirmed_items)} factor{'s' if len(confirmed_items) > 1 else ''} confirmed — {environment}</span>
                  <p style="margin:0.5rem 0 0;font-size:0.95rem;">
                    These factors should be documented and cross-referenced against Section B and
                    Section F of the EHCP.
                  </p>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("---")
                if st.button("Generate EHCP amendment request", key="gen_amendment"):
                    st.session_state.show_amendment = True

                if st.session_state.get("show_amendment"):
                    draft = f"""Dear [SENCO name],

I am writing to request a formal amendment to [child's name]'s Education, Health and Care Plan — specifically to Sections B and F — following observations in the {environment} on {today}.

The following has been identified:

{chr(10).join(f"- {item}" for item in confirmed_items)}

The current Section F does not specify provision for this environment in sufficient detail to be enforceable. I am requesting:

- Specific named provision for the {environment} environment
- Named arrangements confirmed in writing, in place on every occasion
- A half-termly review to confirm adjustments remain in place

Please confirm receipt of this request and advise when it will be considered.

Yours sincerely,
[Your name]
{today}"""

                    st.text_area("Draft amendment request — edit before sending:", value=draft, height=280, key="amendment_text")

                    word_buf = generate_amendment_word(environment, confirmed_items, today)
                    st.download_button(
                        "Download as Word document",
                        data=word_buf,
                        file_name=f"FRED_amendment_request_{today.replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                    if "knowledge_bank" not in st.session_state:
                        st.session_state.knowledge_bank = []
                    entry = {
                        "date": today,
                        "environment": environment,
                        "confirmed": confirmed_items,
                        "amendment_flagged": True,
                    }
                    if entry not in st.session_state.knowledge_bank:
                        st.session_state.knowledge_bank.append(entry)

        if policy_findings:
            st.markdown("---")
            st.markdown("### School policy cross-reference")
            for f in policy_findings:
                tier = f["tier"]
                st.markdown(f"""
                <div class="finding-{tier}">
                  <span class="badge-{tier}">POLICY</span>
                  <p style="font-weight:700;margin:0.4rem 0 0.3rem;">{f['title']}</p>
                  {f'<p style="font-style:italic;color: #555;font-size:0.88rem;margin:0 0 0.4rem;">"{f["extract"]}"</p>' if f.get("extract") else ""}
                  <p style="margin:0;font-size:0.93rem;">{f['commentary']}</p>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### Tone recommendation")

        rec_colour = {"formal": "#C0392B", "collaborative": "#1E8449", "neutral": "#D4A017"}
        rec_bg    = {"formal": "#fdf4f3", "collaborative": "#f3faf5", "neutral": "#fdf9f0"}
        r = tone_rec.get("recommendation", "neutral")

        st.markdown(f"""
        <div style="border-left:4px solid {rec_colour.get(r,'#888')};
                    background:{rec_bg.get(r,'#f9f9f9')};
                    border-radius:0 8px 8px 0;padding:1rem 1.2rem;margin-bottom:1rem;">
          <p style="font-weight:600;margin:0 0 0.3rem;font-size:0.95rem;">
            FRED recommends: {tone_rec.get('label', 'Measured response')}
          </p>
          <p style="margin:0;font-size:0.9rem;line-height:1.6;color: #444;">
            {tone_rec.get('reasoning', '')}
          </p>
          <p style="margin:0.6rem 0 0;font-size:0.78rem;color: #888;">
            Confidence: {tone_rec.get('confidence', '').title()} —
            based on language patterns in this correspondence.
            This is a recommendation, not an instruction.
          </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("**Does this match your experience of the relationship?**")
        tone_confirm = st.radio(
            "",
            ["Yes — use recommended tone",
             "No — it should be more formal",
             "No — it should be more collaborative",
             "I'll decide later"],
            key="tone_confirm",
            label_visibility="collapsed",
            horizontal=True,
        )
        if "formal" in tone_confirm:
            st.session_state.tone_override = "formal"
        elif "collaborative" in tone_confirm:
            st.session_state.tone_override = "collaborative"
        elif "Yes" in tone_confirm:
            st.session_state.tone_override = r

        st.markdown("---")
        st.markdown("### Draft reply")
        st.markdown("""
        <p style='font-size:0.92rem;color: #444;margin-bottom:0.3rem;'>
          Check this is your voice before sending.
        </p>
        <p style='font-size:0.85rem;color: #7ab870;margin-bottom:1rem;font-style:italic;'>
          A focused question in writing is harder to ignore than five.
        </p>
        """, unsafe_allow_html=True)

        st.text_area("Draft reply — edit before sending:", value=draft_reply, height=260, key="draft_reply_display")

        col1, col2 = st.columns(2)
        with col1:
            if draft_reply:
                st.download_button(
                    "Download as Word",
                    data=BytesIO(draft_reply.encode()),
                    file_name=f"FRED_draft_reply_{today.replace(' ','_')}.txt",
                    mime="text/plain",
                )

def page_ehc_journey():
    
def page_ehc_request():
    st.markdown("## Apply for an EHCP")

    st.markdown(f"""
    <div style="background:#eaf5e0;border-radius:10px;padding:1.2rem 1.5rem;margin-bottom:1.5rem;border:0.5px solid #c0ddb0;">
      <p style="margin:0 0 0.6rem;font-size:1rem;color:#2d4a2d;font-weight:600;">
        What this tool does
      </p>
      <p style="margin:0 0 0.5rem;font-size:0.95rem;color:#2d4a2d;line-height:1.7;">
        This tool helps you build a formal request for an Education, Health and Care needs assessment.
        A well-constructed request is more likely to be accepted — and harder for the local authority to refuse.
      </p>
      <p style="margin:0 0 0.5rem;font-size:0.95rem;color:#2d4a2d;line-height:1.7;">
        You will produce a structured written request covering your child's needs across ten categories,
        ready to send to your local authority.
      </p>
      <p style="margin:0;font-size:0.95rem;color:#2d4a2d;line-height:1.7;">
        Your answers save as you go. You can close this page and continue where you left off.
      </p>
    </div>
    """, unsafe_allow_html=True)

    user = st.session_state.get("user")
    existing_request = None

    if SUPABASE_AVAILABLE and supabase and user:
        try:
            supabase.auth.set_session(
                st.session_state["session"].access_token,
                st.session_state["session"].refresh_token
            )
            result = supabase.table("ehc_requests_v2") \
                .select("*") \
                .eq("user_id", str(user.id)) \
                .eq("status", "in_progress") \
                .order("updated_at", desc=True) \
                .limit(1) \
                .execute()
            if result.data:
                existing_request = result.data[0]
        except Exception as e:
            st.caption(f"Could not check for existing request: {e}")

    if existing_request:
        completed = sum(
            1 for i in range(1, 11)
            if existing_request.get(f"category_{i}")
        )
        st.markdown(f"""
        <div style="background:white;border:1px solid #d0dae8;border-radius:8px;
                    padding:1rem 1.2rem;margin-bottom:1.2rem;">
          <p style="margin:0;font-size:0.95rem;">
            You have a request in progress —
            <b>{completed} of 10</b> categories complete.
          </p>
        </div>
        """, unsafe_allow_html=True)

        if st.button("Continue where you left off", use_container_width=False, key="ehc_continue"):
            st.session_state["ehc_request_id"] = existing_request["id"]
            st.session_state["ehc_request_data"] = existing_request
            st.session_state["ehc_request_started"] = True
            st.rerun()

    else:
        if st.button("Start my request", use_container_width=False, key="ehc_start"):
            if SUPABASE_AVAILABLE and supabase and user:
                try:
                    supabase.auth.set_session(
                        st.session_state["session"].access_token,
                        st.session_state["session"].refresh_token
                    )
                    new_row = supabase.table("ehc_requests_v2").insert({
                        "user_id": str(user.id),
                        "status": "in_progress",
                    }).execute()
                    if new_row.data:
                        st.session_state["ehc_request_id"] = new_row.data[0]["id"]
                        st.session_state["ehc_request_data"] = new_row.data[0]
                except Exception as e:
                    st.caption(f"Could not create request: {e}")
            st.session_state["ehc_request_started"] = True
            st.rerun()

    if st.session_state.get("ehc_request_started"):
        if st.session_state.get("ehc_journey_active"):
            page_ehc_journey()
        else:
            st.markdown("---")
            st.markdown("## A few things before you start")
            st.markdown("""
            <div style="background:white;border:1px solid #d0dae8;border-radius:10px;
                        padding:1.4rem 1.6rem;margin-bottom:1.5rem;">
              <p style="margin:0 0 0.9rem;font-size:0.97rem;line-height:1.8;">
                &#9312; &nbsp; Answer in your own words — there is no right or wrong way to describe your child.
              </p>
              <p style="margin:0 0 0.9rem;font-size:0.97rem;line-height:1.8;">
                &#9313; &nbsp; Think about specific moments rather than general descriptions — the more concrete the better.
              </p>
              <p style="margin:0 0 0.9rem;font-size:0.97rem;line-height:1.8;">
                &#9314; &nbsp; You do not need to complete everything in one sitting — your answers save automatically as you go.
              </p>
              <p style="margin:0;font-size:0.97rem;line-height:1.8;">
                &#9315; &nbsp; You know your child better than anyone who will assess them — trust what you know.
              </p>
            </div>
            """, unsafe_allow_html=True)
            if st.button("Let's start", key="ehc_guide_start", use_container_width=False):
                st.session_state["ehc_journey_active"] = True
                st.session_state["ehc_current_category"] = 1
                st.rerun()
                
def page_subscriber():
    st.markdown(f"""
    <div style="background:{NAVY};border-radius:8px;padding:2rem;margin-bottom:1.5rem;">
      <p style="color:#a8b8d8;font-size:0.85rem;margin:0 0 0.3rem;letter-spacing:0.1em;text-transform:uppercase;">
        FRED BETA — SUBSCRIBER WORKSPACE
      </p>
      <h2 style="color:white;margin:0 0 0.4rem;font-family:'Playfair Display',serif;">
        Welcome to FRED
      </h2>
      <p style="color:#c8d8f0;margin:0;font-size:0.95rem;">
        Everything is unlocked. You have full access during beta.
      </p>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2, tab3, tab4 = st.tabs(["📄 My Report", "✉️ Correspondence", "📋 Meeting Prep", "🗂️ Evidence Bank"])

    with tab1:
        st.markdown("### Your EHCP report")
        if st.session_state.findings:
            col1, col2 = st.columns(2)
            with col1:
                word_buf = generate_word_report(
                    st.session_state.findings,
                    doc_type=st.session_state.doc_type,
                    situation=st.session_state.situation
                )
                st.download_button(
                    "Download as Word",
                    data=word_buf,
                    file_name="FRED_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with col2:
                pdf_buf = generate_pdf_report(
                    st.session_state.findings,
                    doc_type=st.session_state.doc_type,
                    situation=st.session_state.situation
                )
                st.download_button(
                    "Download as PDF",
                    data=pdf_buf,
                    file_name="FRED_report.pdf",
                    mime="application/pdf",
                )
            st.markdown("---")
            st.markdown("### All findings")
            for finding in st.session_state.findings:
                render_finding_card(finding, show_full=True)
        else:
            st.info("No report yet. Go back to the home page and upload your EHCP.")
            if st.button("Upload an EHCP"):
                st.session_state.stage = "upload"
                st.rerun()

    with tab2:
        st.markdown("### Correspondence analysis")
        st.markdown(
            "Upload emails from school or the LA. FRED will read them for you — "
            "tone, intent, gaps, and what to do next."
        )
        email1 = st.file_uploader("Most recent email (PDF or Word)", type=["pdf", "docx", "txt"], key="sub_email1")
        email2 = st.file_uploader("Previous email — optional", type=["pdf", "docx", "txt"], key="sub_email2")

        tone_q = st.radio(
            "How is the school engaging right now?",
            ["Genuinely trying to help", "Going through the motions", "Actively avoiding or obstructing"],
            help="You know the relationship. FRED doesn't. Your answer calibrates the analysis."
        )

        if st.button("Analyse correspondence") and email1:
            with st.spinner("Reading correspondence…"):
                text1 = extract_text_from_pdf(email1) if email1.name.endswith(".pdf") else extract_text_from_docx(email1)
                text2 = ""
                if email2:
                    text2 = extract_text_from_pdf(email2) if email2.name.endswith(".pdf") else extract_text_from_docx(email2)
                combined = text1 + "\n\n" + text2
                matched = detect_patterns(combined)
                tone_rec = detect_tone_recommendation(combined, matched)

            st.markdown("### Patterns detected")
            if matched:
                for p in matched[:5]:
                    st.markdown(f"""
                    <div class="finding-{p['tier']}">
                      <span class="badge-{p['tier']}">{p['name']}</span>
                      <p style="margin:0.4rem 0 0.3rem;font-size:0.93rem;">{p['explanation']}</p>
                      <p style="font-style:italic;font-size:0.88rem;color: #555;margin:0;">{p['fred_question']}</p>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.success("No major patterns detected in this correspondence.")

            st.markdown(f"""
            <div style="background:#eaf5e0;border-radius:8px;padding:0.8rem 1.1rem;margin-top:1rem;">
              <p style="font-weight:600;margin:0 0 0.2rem;font-size:0.9rem;">Tone recommendation: {tone_rec['label']}</p>
              <p style="margin:0;font-size:0.85rem;color: #555;">{tone_rec['reasoning']}</p>
            </div>
            """, unsafe_allow_html=True)

    with tab3:
        st.markdown("### Meeting preparation")
        upcoming = st.session_state.get("upcoming_dates", "")
        if upcoming:
            st.info(f"Upcoming date noted: {upcoming}")

        st.markdown("**Before the meeting — check these:**")
        checklist = [
            "Request the agenda in writing at least five working days before",
            "Ask for the delivery log for each provision in Section F",
            "Confirm who will be attending and in what capacity",
            "Take a copy of the EHCP and this report",
            "Note the date, time, and attendees at the start of the meeting",
            "Request that minutes are circulated within five working days",
            "Follow up any verbal agreements in writing within 24 hours",
        ]
        for item in checklist:
            st.checkbox(item, key=f"check_{item[:20]}")

        st.markdown("---")
        st.markdown("**After the meeting:**")
        post = st.text_area(
            "Note anything agreed verbally that needs to be confirmed in writing",
            placeholder="e.g. SENCO agreed to provide weekly session records from next term...",
            height=100,
        )
        if post:
            st.success(
                "Noted. Follow this up in writing within 24 hours. "
                "Subject line: 'Confirmation of agreements — [date] meeting'."
            )

    with tab4:
        st.markdown("### Evidence bank")
        bank = st.session_state.get("knowledge_bank", [])
        held = st.session_state.get("held_findings", [])
        if bank:
            for entry in bank:
                confirmed_list = entry.get("confirmed", [])
                items_html = "".join(f'<p style="font-size:0.85rem;margin:0.1rem 0;">• {item}</p>' for item in confirmed_list)
                amendment_html = "<p style='font-size:0.82rem;font-weight:700;color: #C0392B;margin-top:0.4rem;'>⚑ EHCP amendment flagged</p>" if entry.get("amendment_flagged") else ""
                st.markdown(f"""
                <div style="background:white;border:1px solid #d0dae8;border-radius:6px;padding:1rem 1.2rem;margin-bottom:0.8rem;">
                  <p style="font-weight:700;margin:0 0 0.2rem;">{entry['date']} — {entry['environment']}</p>
                  <p style="margin:0 0 0.4rem;font-size:0.88rem;color: #555;">{len(confirmed_list)} confirmed factor(s)</p>
                  {items_html}
                  {amendment_html}
                </div>
                """, unsafe_allow_html=True)
        elif held:
            for entry in held:
                st.markdown(f"""
                <div style="background:white;border:1px solid #d0dae8;border-radius:6px;padding:1rem 1.2rem;margin-bottom:0.8rem;">
                  <p style="font-weight:700;margin:0 0 0.2rem;">{entry['date']}</p>
                  <p style="margin:0;font-size:0.88rem;color: #555;">{entry.get('patterns',0)} pattern(s) — held without action</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.markdown(
                "<p style='color:#888;font-size:0.9rem;'>"
                "No evidence entries yet."
                "</p>",
                unsafe_allow_html=True
            )

    st.markdown("---")
    st.markdown(f"""
    <a href="{GOOGLE_FORM_URL}" target="_blank" style="
        display:inline-block;
        background:#e8eef8;
        color:{NAVY};
        padding:0.6rem 1.4rem;
        border-radius:4px;
        text-decoration:none;
        font-family:'Source Sans 3',sans-serif;
        font-weight:600;
        font-size:0.9rem;
    ">Leave feedback on FRED →</a>
    """, unsafe_allow_html=True)


# ── NAVIGATION ────────────────────────────────────────────────────────────────

def render_nav():
    # Scroll-to-top on nav: not reliably achievable in Streamlit Community Cloud
    # due to iframe sandboxing. Parked at Session 12. Revisit if self-hosting on Render.
    user = st.session_state.get("user")
    col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 1, 1])
    with col1:
        if st.button("Home", key="nav_home", use_container_width=True):
            st.session_state.stage = "landing"
            st.rerun()
    with col2:
        if st.button("My report", key="nav_report", use_container_width=True):
            if st.session_state.findings:
                st.session_state.stage = "full_report"
                st.rerun()
    with col3:
        if st.button("Correspondence", key="nav_correspondence", use_container_width=True):
            if st.session_state.email_submitted:
                st.session_state.stage = "correspondence"
            else:
                st.session_state.stage = "sneak_peek"
            st.rerun()
    with col4:
        if user:
            st.markdown(
                f"<p style='font-size:0.78rem;color:#666;margin:0.4rem 0 0;text-align:center;'>"
                f"{user.email}</p>",
                unsafe_allow_html=True
            )
    with col5:
        if st.button("Sign out", key="signout_btn", use_container_width=True):
            try:
                if SUPABASE_AVAILABLE:
                    supabase.auth.sign_out()
            except Exception:
                pass
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()


# ── ROUTER ────────────────────────────────────────────────────────────────────
# ── AUTH GATE ─────────────────────────────────────────────────────────────────
if not st.session_state.get("user"):
    page_login()
else:
    render_nav()
    st.markdown("<hr style='margin:0 0 1.5rem;border-color: #d0dae8;'>", unsafe_allow_html=True)
    stage = st.session_state.stage
    if stage == "landing":
        page_landing()
    elif stage == "explainer":
        page_explainer()
    elif stage == "upload":
        page_upload()
    elif stage == "sneak_peek":
        page_sneak_peek()
    elif stage == "full_report":
        page_full_report()
    elif stage == "subscriber":
        page_subscriber()
    elif stage == "correspondence":
        page_correspondence()
    elif stage == "ehc_request":
        page_ehc_request()
    else:
        page_landing()

