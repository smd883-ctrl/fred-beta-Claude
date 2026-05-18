"""
EHCP Parser — Core Extraction Engine
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from .models import (
    EHCPDocument, EHCPSection, ProvisionTable, ProvisionRow,
    ExtractionMethod, SectionStatus, DocumentFamily,
    confidence_level_from_score,
)
from .section_registry import (
    STATUTORY_SECTIONS, PROVISION_SECTIONS,
    match_section_heading, looks_like_contents_page,
    is_contents_entry, looks_like_two_column,
    score_as_provision_table_header, get_family_for_la,
    FAMILY_PATTERNS, detect_need_area_group,
    is_need_area_grouped_format, NEED_AREA_GROUPS,
)

log = logging.getLogger(__name__)

try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class _Page:
    __slots__ = ("number", "text", "lines", "tables_raw", "is_contents")

    def __init__(self, number: int, text: str, tables_raw=None):
        self.number = number
        self.text = text or ""
        self.lines = [l for l in self.text.splitlines() if l.strip()]
        self.tables_raw = tables_raw or []
        self.is_contents = False


class EHCPParser:

    def __init__(self, preferred_families=None, min_section_words=5, debug=False):
        self.preferred_families = preferred_families or list(FAMILY_PATTERNS.keys())
        self.min_section_words = min_section_words
        self.debug = debug

    def parse(self, path, la_hint=None) -> EHCPDocument:
        path = Path(path)
        doc = EHCPDocument()
        doc.source_path = str(path)
        doc.source_filename = path.name

        if not path.exists():
            doc.parse_errors.append(f"File not found: {path}")
            return doc

        suffix = path.suffix.lower()
        doc.file_type = suffix.lstrip(".")

        families = list(self.preferred_families)
        if la_hint:
            la_family = get_family_for_la(la_hint)
            if la_family and la_family in families:
                families.remove(la_family)
                families.insert(0, la_family)
            doc.la_name = la_hint

        try:
            if suffix == ".pdf":
                pages = self._extract_pdf(path, doc)
            elif suffix in (".docx", ".doc"):
                pages = self._extract_docx(path, doc)
            elif suffix in (".txt", ".text"):
                pages = self._extract_text(path, doc)
            else:
                doc.parse_errors.append(f"Unsupported file type: {suffix}")
                return doc
        except Exception as exc:
            doc.parse_errors.append(f"Extraction failed: {exc}")
            return doc

        if not pages:
            doc.parse_errors.append("No text content extracted")
            return doc

        doc.page_count = len(pages)
        all_text = "\n".join(p.text for p in pages)
        doc.raw_text_sample = all_text[:500]

        self._detect_contents_pages(pages, doc)
        self._detect_two_column(pages, doc)
        self._detect_la_from_text(pages, doc)

        page_texts = [p.text for p in pages]
        if is_need_area_grouped_format(page_texts):
            doc.document_family = DocumentFamily.HYBRID
            sections = self._extract_sections_grouped(pages, doc, families)
        else:
            self._detect_document_family(pages, doc, families)
            sections = self._extract_sections(pages, doc, families)

        for sec in sections:
            key = sec.letter or "UNKNOWN"
            doc.sections.setdefault(key, []).append(sec)

        self._classify_section_coverage(doc)
        self._score_overall_confidence(doc)

        return doc

    def _extract_pdf(self, path, doc):
        pages = []
        if HAS_PYMUPDF:
            try:
                pages = self._extract_pdf_pymupdf(path, doc)
                doc.extraction_method = ExtractionMethod.PYMUPDF
            except Exception as exc:
                doc.add_warning("PYMUPDF_FAIL", str(exc), severity="warning")

        if not pages and HAS_PDFPLUMBER:
            try:
                pages = self._extract_pdf_pdfplumber(path, doc)
                doc.extraction_method = ExtractionMethod.PDFPLUMBER
                doc.add_warning("PDF_FALLBACK_PDFPLUMBER",
                                "Fell back to pdfplumber", severity="info")
            except Exception as exc:
                doc.add_warning("PDF_PLUMBER_FAIL", str(exc), severity="error")

        if not pages:
            doc.parse_errors.append("All PDF extraction methods failed")
            return []

        word_counts = [len(p.text.split()) for p in pages]
        low_text = sum(1 for wc in word_counts if wc < 20)
        if len(pages) > 0 and low_text / len(pages) > 0.70:
            doc.is_scanned = True
            doc.add_warning("SCANNED_PDF",
                            f"{low_text}/{len(pages)} pages have very little text",
                            severity="error")
        return pages

    def _extract_pdf_pymupdf(self, path, doc):
        pages = []
        pdf = fitz.open(str(path))
        for i, page in enumerate(pdf, 1):
            text = page.get_text("text") or ""
            pages.append(_Page(number=i, text=text))
        pdf.close()
        return pages

    def _extract_pdf_pdfplumber(self, path, doc):
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                tables_raw = []
                try:
                    tables_raw = page.extract_tables() or []
                except Exception:
                    pass
                pages.append(_Page(number=i, text=text, tables_raw=tables_raw))
        return pages

    def _extract_docx(self, path, doc):
        if not HAS_DOCX:
            doc.parse_errors.append("python-docx not installed")
            return []
        doc.extraction_method = ExtractionMethod.DOCX
        docx = DocxDocument(str(path))
        current_page = []
        page_texts = []
        for para in docx.paragraphs:
            text = para.text
            if "\f" in text:
                parts = text.split("\f")
                current_page.append(parts[0])
                page_texts.append("\n".join(current_page))
                for part in parts[1:-1]:
                    page_texts.append(part)
                current_page = [parts[-1]]
            else:
                current_page.append(text)
        if current_page:
            page_texts.append("\n".join(current_page))
        if not page_texts:
            all_text = "\n".join(p.text for p in docx.paragraphs)
            page_texts = [all_text]
        return [_Page(number=i, text=t) for i, t in enumerate(page_texts, 1)]

    def _extract_text(self, path, doc):
        doc.extraction_method = ExtractionMethod.PLAINTEXT
        text = path.read_text(encoding="utf-8", errors="replace")
        raw_pages = text.split("\f") if "\f" in text else [text]
        return [_Page(number=i, text=p) for i, p in enumerate(raw_pages, 1)]

    def _detect_contents_pages(self, pages, doc):
        for page in pages:
            if looks_like_contents_page(page.lines):
                page.is_contents = True
                doc.has_contents_page = True
                doc.contents_page_numbers.append(page.number)

    def _detect_two_column(self, pages, doc):
        checked = 0
        for page in pages:
            if page.is_contents:
                continue
            if looks_like_two_column(page.text):
                doc.is_two_column = True
                doc.add_warning("TWO_COLUMN_LAYOUT",
                                f"Two-column layout on page {page.number}",
                                page=page.number, severity="warning")
                break
            checked += 1
            if checked >= 5:
                break

    def _detect_document_family(self, pages, doc, families):
        family_hits = {}
        checked = 0
        for page in pages:
            if page.is_contents:
                continue
            for line in page.lines:
                result = match_section_heading(line, families)
                if result:
                    for fam in families:
                        for pat, _ in _get_compiled_safe(fam):
                            if pat.match(line.strip()):
                                family_hits[fam] = family_hits.get(fam, 0) + 1
                                break
            checked += 1
            if checked >= 10:
                break

        if family_hits:
            dominant = max(family_hits, key=family_hits.get)
            family_map = {
                "standard_lettered": DocumentFamily.STANDARD_LETTERED,
                "alpha_dash": DocumentFamily.ALPHA_DASH,
                "numeric_part": DocumentFamily.NUMERIC_PART,
                "bold_heading_only": DocumentFamily.BOLD_HEADING_ONLY,
            }
            doc.document_family = family_map.get(dominant, DocumentFamily.HYBRID)
            if len(family_hits) > 1:
                doc.document_family = DocumentFamily.HYBRID
                doc.add_warning("HYBRID_FORMAT",
                                f"Multiple format families: {list(family_hits.keys())}",
                                severity="info")

    def _detect_la_from_text(self, pages, doc):
        if doc.la_name:
            return
        LA_RE = re.compile(
            r"(?:London Borough of|City of|County Council|Borough Council|"
            r"District Council|Metropolitan Borough|Unitary Authority)[^\n]{0,60}",
            re.IGNORECASE
        )
        sample = pages[:2] + (pages[-1:] if len(pages) > 2 else [])
        for page in sample:
            m = LA_RE.search(page.text)
            if m:
                doc.la_name = m.group(0).strip()
                break

    def _build_line_list(self, pages):
        all_lines = []
        for page in pages:
            for line in page.lines:
                if page.is_contents and is_contents_entry(line):
                    continue
                all_lines.append((page.number, line))
        return all_lines

    def _extract_sections(self, pages, doc, families):
        all_lines = self._build_line_list(pages)

        heading_positions = []
        for idx, (page_num, line) in enumerate(all_lines):
            result = match_section_heading(line, families)
            if result:
                letter, title, conf = result
                heading_positions.append((idx, page_num, letter, title, conf))

        if not heading_positions:
            doc.add_warning("NO_SECTION_HEADINGS",
                            "No section headings found",
                            severity="error")
            return []

        sections = []
        for i, (start_idx, start_page, letter, title, conf) in enumerate(heading_positions):
            body_start = start_idx + 1
            if i + 1 < len(heading_positions):
                body_end = heading_positions[i + 1][0]
                end_page = heading_positions[i + 1][1]
            else:
                body_end = len(all_lines)
                end_page = pages[-1].number

            body_lines = [all_lines[j][1] for j in range(body_start, body_end)]
            body_text = "\n".join(body_lines).strip()
            word_count = len(body_text.split())

            sec = EHCPSection(
                letter=letter,
                label=letter,
                title=title or self._infer_title(letter),
                body_text=body_text if word_count >= self.min_section_words else None,
                page_start=start_page,
                page_end=end_page,
                confidence=conf,
                confidence_level=confidence_level_from_score(conf),
            )

            if word_count < self.min_section_words:
                sec.status = SectionStatus.EMPTY
                sec.extraction_notes.append(f"Only {word_count} words in body")
            else:
                sec.status = SectionStatus.FOUND

            if self._body_looks_like_toc(body_text):
                sec.status = SectionStatus.CONTENTS_ONLY
                sec.from_contents_page = True
                doc.add_warning("TOC_FALSE_POSITIVE",
                                f"Section {letter} body looks like TOC",
                                section=letter, page=start_page)

            sections.append(sec)

        self._flag_duplicate_sections(sections, doc)
        self._extract_provision_tables(sections, pages, doc)
        return sections

    def _extract_sections_grouped(self, pages, doc, families):
        all_lines = self._build_line_list(pages)

        group_positions = []
        heading_positions = []

        for idx, (page_num, line) in enumerate(all_lines):
            group = detect_need_area_group(line)
            if group:
                group_positions.append((idx, page_num, group))
                continue
            result = match_section_heading(line, families)
            if result:
                letter, title, conf = result
                heading_positions.append((idx, page_num, letter, title, conf))

        if not group_positions:
            doc.add_warning("GROUPED_FORMAT_NOT_DETECTED",
                            "Need-area groups not found — using standard extraction",
                            severity="info")
            return self._extract_sections(pages, doc, families)

        doc.add_warning("NEED_AREA_GROUPED_FORMAT",
                        f"{len(group_positions)} need-area groups detected "
                        "(Warwickshire/Capita Synergy format)",
                        severity="info")

        sections = []

        for g_idx, (g_start, g_page, group) in enumerate(group_positions):
            if g_idx + 1 < len(group_positions):
                g_end = group_positions[g_idx + 1][0]
            else:
                g_end = len(all_lines)

            group_headings = [
                h for h in heading_positions
                if g_start < h[0] < g_end
            ]

            if not group_headings:
                continue

            for h_idx, (h_start, h_page, letter, title, conf) in enumerate(group_headings):
                body_start = h_start + 1
                if h_idx + 1 < len(group_headings):
                    body_end = group_headings[h_idx + 1][0]
                    end_page = group_headings[h_idx + 1][1]
                else:
                    body_end = g_end
                    end_page = g_page

                body_lines = [all_lines[j][1] for j in range(body_start, body_end)]
                body_text = "\n".join(body_lines).strip()
                word_count = len(body_text.split())

                sec = EHCPSection(
                    letter=letter,
                    label=letter + " (" + group["label"] + ")",
                    title=title or self._infer_title(letter),
                    body_text=body_text if word_count >= self.min_section_words else None,
                    page_start=h_page,
                    page_end=end_page,
                    confidence=conf,
                    confidence_level=confidence_level_from_score(conf),
                )
                sec.status = (SectionStatus.FOUND
                              if word_count >= self.min_section_words
                              else SectionStatus.EMPTY)
                sec.extraction_notes.append("Need area: " + group["label"])
                sections.append(sec)

        self._flag_duplicate_sections(sections, doc)
        self._extract_provision_tables(sections, pages, doc)
        return sections

    def _body_looks_like_toc(self, body_text):
        lines = [l.strip() for l in body_text.splitlines() if l.strip()]
        if len(lines) < 3:
            return False
        toc_lines = sum(
            1 for l in lines
            if re.search(r'\b\d+\s*$', l) and is_contents_entry(l)
        )
        return toc_lines >= len(lines) * 0.5

    def _flag_duplicate_sections(self, sections, doc):
        letter_count = {}
        for sec in sections:
            if sec.letter:
                letter_count[sec.letter] = letter_count.get(sec.letter, 0) + 1

        for letter, count in letter_count.items():
            if count > 1:
                if letter in ("F", "B", "E"):
                    doc.add_warning("MULTIPLE_SECTION_" + letter,
                                    f"Section {letter} appears {count} times — "
                                    "sub-sections preserved",
                                    section=letter, severity="info")
                    same = [s for s in sections if s.letter == letter]
                    for idx, s in enumerate(same, 1):
                        if s.label == letter:
                            s.label = letter + str(idx)
                else:
                    doc.add_warning("DUPLICATE_SECTION_" + letter,
                                    f"Section {letter} appears {count} times",
                                    section=letter, severity="warning")

    def _extract_provision_tables(self, sections, pages, doc):
        provision_secs = [s for s in sections if s.letter in PROVISION_SECTIONS]
        if not provision_secs:
            return
        page_lookup = {p.number: p for p in pages}
        for sec in provision_secs:
            if not sec.page_start or not sec.page_end:
                continue
            for pnum in range(sec.page_start, sec.page_end + 1):
                page = page_lookup.get(pnum)
                if not page or not page.tables_raw:
                    continue
                for raw_table in page.tables_raw:
                    pt = self._parse_provision_table(raw_table, pnum, sec.letter)
                    if pt is not None:
                        sec.tables.append(pt)

    def _parse_provision_table(self, raw_table, page_num, section_letter):
        if not raw_table or len(raw_table) < 2:
            return None

        def clean(cell):
            if cell is None:
                return ""
            return str(cell).strip().replace("\n", " ")

        rows = [[clean(c) for c in row] for row in raw_table]
        header_row = rows[0]
        score = score_as_provision_table_header(header_row)
        min_score = 0.20 if section_letter == "F" else 0.45
        if score < min_score:
            return None

        pt = ProvisionTable(
            page_number=page_num,
            headers=header_row,
            column_count=len(header_row),
            extraction_method=ExtractionMethod.PDFPLUMBER,
            confidence=score,
        )

        for data_row in rows[1:]:
            if not any(data_row):
                continue
            pr = ProvisionRow(raw_cells=data_row)
            pr = self._assign_provision_row_fields(pr, header_row, data_row)
            pt.rows.append(pr)

        return pt if pt.rows else None

    def _assign_provision_row_fields(self, pr, headers, cells):
        PROVISION_COL = re.compile(r"provision|support|intervention", re.I)
        FREQ_COL = re.compile(r"frequenc|how\s+often|times", re.I)
        WHO_COL = re.compile(r"who|deliver|responsib|staff", re.I)
        SETTING_COL = re.compile(r"where|setting|location", re.I)

        for i, header in enumerate(headers):
            if i >= len(cells):
                break
            val = cells[i]
            if not val:
                continue
            if PROVISION_COL.search(header):
                pr.provision_text = val
            elif FREQ_COL.search(header):
                pr.frequency = val
            elif WHO_COL.search(header):
                pr.who_delivers = val
            elif SETTING_COL.search(header):
                pr.setting = val

        if not pr.provision_text:
            non_empty = [c for c in cells if c]
            if non_empty:
                pr.provision_text = non_empty[0]

        time_re = re.compile(r'\b\d+\s*(hour|min|session|week|term)', re.I)
        provision_text = pr.provision_text or " ".join(cells)
        pr.quantified = bool(time_re.search(provision_text))
        return pr

    def _classify_section_coverage(self, doc):
        found = set(doc.sections.keys())
        for letter in STATUTORY_SECTIONS:
            l_upper = letter.upper()
            if l_upper in found:
                secs = doc.sections[l_upper]
                if all(s.status == SectionStatus.CONTENTS_ONLY for s in secs):
                    doc.sections_ambiguous.append(l_upper)
                elif any(s.status == SectionStatus.AMBIGUOUS for s in secs):
                    doc.sections_ambiguous.append(l_upper)
                else:
                    doc.sections_found.append(l_upper)
            else:
                doc.sections_missing.append(l_upper)

    def _score_overall_confidence(self, doc):
        score = 1.0
        if doc.is_scanned:
            score -= 0.50
        score -= len(doc.parse_errors) * 0.15
        if "F" in doc.sections_missing:
            score -= 0.25
        if doc.is_two_column:
            score -= 0.10
        score -= len(doc.sections_ambiguous) * 0.05

        all_sections = [s for secs in doc.sections.values() for s in secs]
        if all_sections:
            avg_sec_conf = sum(s.confidence for s in all_sections) / len(all_sections)
            score = score * 0.6 + avg_sec_conf * 0.4

        if not doc.sections_found:
            score = min(score, 0.15)

        doc.overall_confidence = round(max(0.0, min(score, 1.0)), 3)
        doc.overall_confidence_level = confidence_level_from_score(doc.overall_confidence)

    @staticmethod
    def _infer_title(letter):
        TITLES = {
            "A": "Views, Interests and Aspirations",
            "B": "Special Educational Needs",
            "C": "Health Needs",
            "D": "Social Care Needs",
            "E": "Educational Outcomes",
            "F": "Special Educational Provision",
            "G": "Health Provision",
            "H1": "Social Care Provision",
            "H2": "Other Social Care Provision",
            "I": "Educational Placement",
            "J": "Personal Budget",
            "K": "Appendices",
        }
        return TITLES.get(letter.upper(), f"Section {letter}")


def _get_compiled_safe(family):
    from .section_registry import _get_compiled
    return _get_compiled(family)


def parse_ehcp(path, la_hint=None, debug=False) -> EHCPDocument:
    parser = EHCPParser(debug=debug)
    return parser.parse(path, la_hint=la_hint)
